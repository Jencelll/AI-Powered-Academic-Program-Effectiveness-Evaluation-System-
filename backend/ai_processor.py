# ai_processor.py (Updated to capture ALL analysis output)
import os, re, warnings
import pandas as pd
import numpy as np
from datetime import datetime, timezone, timedelta
from collections import Counter
from docx import Document
from difflib import get_close_matches
from sklearn.cluster import KMeans
from sklearn.metrics import precision_score, recall_score, f1_score, accuracy_score
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import random
import json # For storing complex data like analytics summaries
from models import db, FacultyUpload, SubjectAnalysis, AnalyticsData
import io
from contextlib import redirect_stdout, redirect_stderr # Import for capturing print output
import plotly.express as px # Import for the subject category bar chart
import pdfplumber # For PDF extraction

warnings.filterwarnings('ignore')

# Paths (Adjust these for your web system)
BASE_DIR = os.path.abspath(os.path.dirname(__file__)) # Use project root
FACULTY_DIR = os.path.join(BASE_DIR, 'static', 'uploads')
RESULTS_DIR = os.path.join(BASE_DIR, 'results') # Or store in DB directly
MASTER_PATH = os.path.join(BASE_DIR, 'historical_data', 'CQI_Master.csv')
MODEL_PATH = os.path.join(BASE_DIR, 'ai_models')
LSPU_DATA_PATH = os.path.join(BASE_DIR, 'LSPU_Assessment_Data.xlsx') # Path to the uploaded LSPU data

os.makedirs(FACULTY_DIR, exist_ok=True)
os.makedirs(RESULTS_DIR, exist_ok=True)
os.makedirs(os.path.dirname(MASTER_PATH), exist_ok=True)
os.makedirs(MODEL_PATH, exist_ok=True)

# Global variable to store students from the class record (for top performers within that specific record)
CLASS_RECORD_STUDENTS = pd.DataFrame()

# --- Utility functions ---
def clean_str(x):
    return re.sub(r'\s+', ' ', str(x)).strip() if pd.notna(x) and str(x) != 'nan' else ""

def is_valid_full_name(value: str):
    s = clean_str(value)
    if not s:
        return False
    up = s.upper()
    bad = (
        'STUDENT', 'TOTAL', 'AVERAGE', 'SUMMARY', 'SCHEDULE', 'PREPARED',
        'CERTIFIED', 'CHECKED', 'SUBMITTED', 'INSTRUCTOR', 'DEPARTMENT',
        'NOTED', 'NOTE'
    )
    if any(k in up for k in bad):
        return False
    if ':' in s:
        return False
    try:
        if re.search(r"\d", s):
            return False
        words = re.split(r"[,\s]+", s)
        words = [w for w in words if w]
        if len(words) < 2:
            return False
        letters = sum(ch.isalpha() for ch in s)
        if letters < 4:
            return False
        return True
    except Exception:
        return False

def make_unique_columns(cols):
    seen = {}
    out = []
    for c in cols:
        key = clean_str(c).upper()
        if key in seen:
            seen[key] += 1
            out.append(f"{key}__{seen[key]}")
        else:
            seen[key] = 0
            out.append(key)
    return out

def sanitize_series(sr):
    return pd.to_numeric(sr.astype(str).str.replace(r'[^0-9.-]', '', regex=True), errors='coerce').fillna(0)

def _find_header_index(cols, words):
    uw = [w.upper() for w in words]
    for i, c in enumerate(cols):
        uc = str(c).upper()
        if all(w in uc for w in uw):
            return i
    return -1

def find_group_columns_by_header(df, group_kw):
    cols = list(df.columns)
    i_m = _find_header_index(cols, [group_kw, 'MALE'])
    i_f = _find_header_index(cols, [group_kw, 'FEMALE'])
    i_t = _find_header_index(cols, [group_kw, 'TOTAL'])
    if i_m >= 0 and i_f >= 0 and i_t >= 0:
        return {'male_idx': i_m, 'female_idx': i_f, 'total_idx': i_t}
    # Try proximity around total column
    if i_t >= 0:
        for offset in range(1, 4):
            left = i_t - offset
            right = i_t + offset
            if left >= 0 and 'MALE' in str(cols[left]).upper():
                i_m = left
            if left >= 0 and 'FEMALE' in str(cols[left]).upper():
                i_f = left
            if right < len(cols) and 'MALE' in str(cols[right]).upper():
                i_m = right
            if right < len(cols) and 'FEMALE' in str(cols[right]).upper():
                i_f = right
            if i_m >= 0 and i_f >= 0:
                return {'male_idx': i_m, 'female_idx': i_f, 'total_idx': i_t}
    return None

def process_subject_accomplishment(filepath: str, file_type: str, upload_id: int):
    """
    Parses Subject Accomplishment Report (PDF, DOCX, XLSX).
    Extracts: Subject Code/Name, Weakness, Action Taken, Recommendation.
    Saves to SubjectAccomplishment table.
    """
    from models import SubjectAccomplishment, SubjectAnalysis
    
    extracted_data = [] # List of dicts {subject, weakness, action, recommendation}
    
    filename = os.path.basename(filepath)
    subject_from_filename = None
    # Try to extract subject from filename if possible (e.g. "Accomplishment_ITST301.pdf")
    # Pattern: [A-Z]{3,4}[\s_-]?\d{3}
    match = re.search(r'([A-Z]{3,4}[\s_-]?\d{3})', filename, re.IGNORECASE)
    if match:
        subject_from_filename = match.group(1).upper().replace('_', ' ').replace('-', ' ')

    print(f"📄 Processing Subject Accomplishment: {filename} ({file_type})")
    
    try:
        if file_type == 'pdf':
            extracted_data = _parse_accomplishment_pdf(filepath)
        elif file_type in ['xlsx', 'xls']:
            extracted_data = _parse_accomplishment_excel(filepath)
        elif file_type == 'docx':
            extracted_data = _parse_accomplishment_docx(filepath)
        else:
            print("❌ Unsupported file type for accomplishment report.")
            return

        if not extracted_data:
            print("⚠️ No data extracted from accomplishment report.")
            return

        # Save to DB
        count = 0
        for item in extracted_data:
            # Determine subject: extracted from doc > extracted from filename
            subj = item.get('subject') or subject_from_filename
            if not subj:
                # If still no subject, maybe skip or mark as 'Unknown'
                subj = "UNKNOWN"
            
            # clean subject string
            subj = clean_str(subj).upper()
            
            # Try to link to SubjectAnalysis
            analysis_rec = None
            # Fetch all SubjectAnalysis to find a match
            # This is a simple containment match: if subj (e.g. ITST301) is in analysis.course (e.g. ITST 301 - Advanced Database)
            all_analyses = SubjectAnalysis.query.all()
            for analysis in all_analyses:
                # Normalize course name from DB
                course_norm = clean_str(analysis.course).upper().replace(' ', '').replace('-', '').replace('_', '')
                subj_norm = subj.replace(' ', '').replace('-', '').replace('_', '')
                
                if subj_norm in course_norm or course_norm in subj_norm:
                    analysis_rec = analysis
                    break
            
            # Check for duplicates before adding
            existing_dup = SubjectAccomplishment.query.filter_by(
                subject_code=subj,
                weakness=item.get('weakness'),
                action_taken=item.get('action'),
                recommendation=item.get('recommendation')
            ).first()
            
            if existing_dup:
                print(f"  ⚠️ Skipping duplicate accomplishment for {subj}")
                continue

            rec = SubjectAccomplishment(
                subject_code=subj,
                subject_analysis_id=analysis_rec.id if analysis_rec else None,
                weakness=item.get('weakness'),
                action_taken=item.get('action'),
                recommendation=item.get('recommendation'),
                upload_id=upload_id
            )
            db.session.add(rec)
            count += 1
        
        db.session.commit()
        print(f"✅ Saved {count} accomplishment records.")

    except Exception as e:
        print(f"❌ Error processing accomplishment report: {str(e)}")
        import traceback
        traceback.print_exc()

def _parse_accomplishment_pdf(filepath):
    results = []
    subject_found = None
    
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            # Try to find subject in text header if not found yet
            if not subject_found and text:
                # Look for "REPORT in <Subject>" or "Subject: <Subject>"
                # Pattern for screenshot: "ACADEMIC ACCOMPLISHMENT REPORT in ITST 301"
                m = re.search(r'(?:REPORT\s+in|SUBJECT[:\s])\s+([A-Z0-9\s-]+(?:–|-)\s+[A-Za-z\s]+)', text, re.IGNORECASE)
                if not m:
                     # simpler pattern: ITST 301
                     m = re.search(r'\b([A-Z]{3,4}\s?\d{3})\b', text)
                if m:
                    subject_found = m.group(1).strip()
            
            # Extract tables
            tables = page.extract_tables()
            for table in tables:
                # Find headers
                if not table: continue
                headers = [str(h).upper().replace('\n', ' ').strip() for h in table[0] if h]
                
                # Map columns
                idx_weak = -1
                idx_action = -1
                idx_rec = -1
                
                for i, h in enumerate(headers):
                    if 'WEAKNESS' in h or 'PROBLEM' in h: idx_weak = i
                    elif 'ACTION' in h: idx_action = i
                    elif 'RECOMMENDATION' in h: idx_rec = i
                
                if idx_weak != -1 and idx_action != -1 and idx_rec != -1:
                    # Valid table found, iterate rows
                    for row in table[1:]:
                        if not row: continue
                        # Ensure row has enough columns
                        if len(row) <= max(idx_weak, idx_action, idx_rec): continue
                        
                        w = clean_str(row[idx_weak])
                        a = clean_str(row[idx_action])
                        r = clean_str(row[idx_rec])
                        
                        if w or a or r:
                            results.append({
                                'subject': subject_found,
                                'weakness': w,
                                'action': a,
                                'recommendation': r
                            })
    return results

def _parse_accomplishment_excel(filepath):
    results = []
    xls = pd.ExcelFile(filepath)
    for sheet in xls.sheet_names:
        df = xls.parse(sheet)
        # Scan for header row
        # We look for a row that contains "WEAKNESS" and "ACTION"
        header_idx = -1
        for i, row in df.iterrows():
            row_str = row.astype(str).str.upper().str.cat(sep=' ')
            if 'WEAKNESS' in row_str and 'ACTION' in row_str:
                header_idx = i
                break
        
        if header_idx != -1:
            # Reload with header
            df = xls.parse(sheet, header=header_idx+1)
            # Find cols
            cols = [str(c).upper().strip() for c in df.columns]
            col_weak = next((c for c in cols if 'WEAKNESS' in c or 'PROBLEM' in c), None)
            col_action = next((c for c in cols if 'ACTION' in c), None)
            col_rec = next((c for c in cols if 'RECOMMENDATION' in c), None)
            
            if col_weak and col_action and col_rec:
                for _, row in df.iterrows():
                    w = clean_str(row.get(df.columns[cols.index(col_weak)]))
                    a = clean_str(row.get(df.columns[cols.index(col_action)]))
                    r = clean_str(row.get(df.columns[cols.index(col_rec)]))
                    
                    if w or a or r:
                        results.append({
                            'subject': None, # Subject hard to extract from excel cells reliably without specific cell address, rely on filename or pass logic later
                            'weakness': w,
                            'action': a,
                            'recommendation': r
                        })
    return results

def _parse_accomplishment_docx(filepath):
    results = []
    doc = Document(filepath)
    subject_found = None
    
    # Try text search for subject
    full_text = '\n'.join([p.text for p in doc.paragraphs])
    m = re.search(r'(?:REPORT\s+in|SUBJECT[:\s])\s+([A-Z0-9\s-]+(?:–|-)\s+[A-Za-z\s]+)', full_text, re.IGNORECASE)
    if not m:
        m = re.search(r'\b([A-Z]{3,4}\s?\d{3})\b', full_text)
    if m:
        subject_found = m.group(1).strip()

    for table in doc.tables:
        # Check header row
        headers = [c.text.upper().strip().replace('\n', ' ') for c in table.rows[0].cells]
        
        idx_weak = -1
        idx_action = -1
        idx_rec = -1
        
        for i, h in enumerate(headers):
            if 'WEAKNESS' in h or 'PROBLEM' in h: idx_weak = i
            elif 'ACTION' in h: idx_action = i
            elif 'RECOMMENDATION' in h: idx_rec = i
            
        if idx_weak != -1 and idx_action != -1 and idx_rec != -1:
             for row in table.rows[1:]:
                cells = row.cells
                if len(cells) <= max(idx_weak, idx_action, idx_rec): continue
                
                w = clean_str(cells[idx_weak].text)
                a = clean_str(cells[idx_action].text)
                r = clean_str(cells[idx_rec].text)
                
                if w or a or r:
                    results.append({
                        'subject': subject_found,
                        'weakness': w,
                        'action': a,
                        'recommendation': r
                    })
    return results

def _infer_group_triplets(df, start_idx=0):
    cols = list(df.columns)
    n = len(cols)
    def numcol(i):
        try:
            s = pd.to_numeric(df.iloc[:, i], errors='coerce')
            return s
        except Exception:
            return pd.Series([np.nan]*len(df))
    groups = []
    i = max(0, start_idx)
    while i + 2 < n and len(groups) < 5:
        a = numcol(i)
        b = numcol(i+1)
        c = numcol(i+2)
        valid = (~a.isna()) & (~b.isna()) & (~c.isna())
        if valid.sum() >= max(3, int(0.4*len(df))):
            ab = (a.fillna(0) + b.fillna(0))
            diff = (c.fillna(0) - ab).abs()
            # Accept if median difference is small
            med = float(diff[valid].median() if valid.any() else 0)
            if med <= 1.0:
                groups.append({'male_idx': i, 'female_idx': i+1, 'total_idx': i+2})
                i += 3
                continue
        i += 1
    return groups

def smart_find_col(df, possible_names):
    """Flexible column finder using substring and fuzzy matching."""
    if df.empty:
        return None
    df_cols = [str(c).strip() for c in df.columns]
    
    # 1. Exact case-insensitive match
    for name in possible_names:
        for col in df_cols:
            if name.lower() == col.lower():
                return col

    # 2. Substring match (name in col)
    for name in possible_names:
        for col in df_cols:
            if name.lower() in col.lower():
                 return col
                 
    # 3. Fuzzy match with higher cutoff
    for name in possible_names:
        close = get_close_matches(name.lower(), [c.lower() for c in df_cols], n=1, cutoff=0.85)
        if close:
            return next(c for c in df_cols if c.lower() == close[0])
            
    return None

def preview_excel(path, n=5):
    print(f"--- Preview of Spreadsheet: {os.path.basename(path)} ---")
    try:
        if str(path).lower().endswith('.csv'):
            df = pd.read_csv(path, nrows=n)
        else:
            df = pd.read_excel(path, nrows=n)
        print(df.head(n)) # Use print instead of display for backend
        return df
    except Exception as e:
        print(f"❌ Could not preview spreadsheet: {e}")
        return None

def preview_docx(path, n=5):
    print(f"--- Preview of DOCX: {os.path.basename(path)} ---")
    try:
        doc = Document(path)
        tables = []
        for table in doc.tables:
            data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            # Ensure table has header and rows before creating DataFrame
            if data and len(data) > 1 and len(data[0]) > 0:
                try:
                    df = pd.DataFrame(data[1:], columns=data[0])
                except Exception:
                    # Fallback: no header alignment, use raw data
                    df = pd.DataFrame(data)
                print(df.head(n)) # Use print instead of display for backend
                return df
        print("No tables found. Showing first paragraphs:")
        for i, p in enumerate(doc.paragraphs[:n]):
            print(f"{i+1}. {clean_str(p.text)}")
        return None
    except Exception as e:
        print(f"❌ DOCX preview failed: {e}")
        return None

def preview_pdf(path, n=5):
    print(f"--- Preview of PDF: {os.path.basename(path)} ---")
    tables_collected = []
    text_fragments = []
    try:
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                if getattr(pdf, 'is_encrypted', False):
                    print("⚠️ PDF appears to be encrypted; skipping.")
                    return None
                page_count = len(pdf.pages)
                print(f"   Pages: {page_count}")
                for idx, page in enumerate(pdf.pages[:min(page_count, 10)]):
                    try:
                        txt = page.extract_text() or ''
                        if txt.strip():
                            text_fragments.append(txt.strip())
                        # Try table extraction where possible
                        try:
                            tables = page.extract_tables() or []
                            for t in tables:
                                if t and any(any(c and str(c).strip() for c in row) for row in t):
                                    tables_collected.append(t)
                        except Exception as te:
                            print(f"   ⚠️ Table extraction error on page {idx+1}: {te}")
                    except Exception as pe:
                        print(f"   ⚠️ Page {idx+1} extraction error: {pe}")
        except ImportError:
            print("pdfplumber not available; falling back to PyPDF2.")
            import PyPDF2
            with open(path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                if reader.is_encrypted:
                    print("⚠️ PDF is encrypted; skipping.")
                    return None
                page_count = len(reader.pages)
                print(f"   Pages: {page_count}")
                for i in range(min(page_count, 10)):
                    try:
                        txt = reader.pages[i].extract_text() or ''
                        if txt.strip():
                            text_fragments.append(txt.strip())
                    except Exception as pe:
                        print(f"   ⚠️ Page {i+1} extraction error: {pe}")

        # OCR fallback if no text and no tables
        if not text_fragments and not tables_collected:
            try:
                import pypdfium2 as pdfium
                import pytesseract
                from PIL import Image
                pdf = pdfium.PdfDocument(path)
                pc = len(pdf)
                print(f"   OCR fallback: {pc} pages")
                for i in range(min(pc, 3)):
                    page = pdf[i]
                    pil = page.render(scale=2).to_pil()
                    pil = pil.convert('L')
                    text = pytesseract.image_to_string(pil)
                    if text.strip():
                        text_fragments.append(text.strip())
            except Exception as ocr_err:
                print(f"   ⚠️ OCR fallback failed: {ocr_err}")

        # Try to normalize first detected table to DataFrame
        if tables_collected:
            import pandas as pd
            for tbl in tables_collected:
                # Use first non-empty table
                rows = [[str(c).strip() if c is not None else '' for c in row] for row in tbl]
                if not rows:
                    continue
                # Combine first two header rows if present to retain group labels
                if len(rows) >= 2:
                    r0 = rows[0]
                    r1 = rows[1]
                    m = max(len(r0), len(r1))
                    header = []
                    for i in range(m):
                        h0 = (r0[i] if i < len(r0) else '').strip()
                        h1 = (r1[i] if i < len(r1) else '').strip()
                        combined = ' '.join([h0, h1]).strip()
                        header.append(combined if combined else (h0 or h1 or f'COL_{i+1}'))
                    data = rows[2:]
                else:
                    header = rows[0]
                    data = rows[1:] if len(rows) > 1 else []
                try:
                    df = pd.DataFrame(data, columns=header)
                except Exception:
                    df = pd.DataFrame(rows)
                print(df.head(n))
                return df

        # If no tables, print text preview for debugging
        joined = '\n'.join(text_fragments)
        if not joined.strip():
            print("⚠️ No extractable text found. The PDF may be image-only.")
            return None
        print(joined[:1200])
        return None
    except Exception as e:
        print(f"❌ PDF preview failed: {e}")
        import traceback
        traceback.print_exc()
        return None

def load_master():
    if os.path.exists(MASTER_PATH):
        return pd.read_csv(MASTER_PATH)
    return pd.DataFrame()

def save_master(df):
    df.to_csv(MASTER_PATH, index=False)
    print(f"📚 Master dataset saved → {MASTER_PATH}")

# ------------------------------------------------------------
# NEW: Process Class Record File - Identifies top performers from the class record itself
# ------------------------------------------------------------
def process_class_record_students(class_record_path, course_code_from_profile, section_from_profile):
    """Reads the class record Excel file and identifies top performers based on grade."""
    if not class_record_path:
        return pd.DataFrame()

    print(f"📄 Preview: Class Record ({os.path.basename(class_record_path)}) - Looking for section: {section_from_profile} (for course {course_code_from_profile})")
    try:
        # Read the 'Summary' sheet (or first sheet if 'Summary' doesn't exist)
        # This is the key change from the original code.
        try:
            df_sheet = pd.read_excel(class_record_path, sheet_name='Summary', header=None)
        except ValueError:
            # If 'Summary' sheet doesn't exist, try the first sheet
            xls = pd.ExcelFile(class_record_path)
            sheet_name = xls.sheet_names[0]
            df_sheet = pd.read_excel(class_record_path, sheet_name=sheet_name, header=None)
            print(f"   Using first sheet: {sheet_name}")

        print(f"   Processing sheet: Summary (or first available sheet)")

        # Find the header row for the student data table.
        # The header row is at index 12 in your image, containing: 'RANK', 'STUDENTS' NAME', 'MIDTERM', 'FINALS', 'FINAL GRADE', etc.
        header_row = None
        for i in range(min(25, df_sheet.shape[0])):  # Check first 25 rows
            row = df_sheet.iloc[i].astype(str).str.upper()
            row_text = ' '.join(row.values).upper()
            # Look for the exact column headers or close matches
            if any(word in row_text for word in ['STUDENTS\' NAME', 'GRADE', 'RANK', 'FINAL']):
                # More precise check: look for the sequence of expected headers
                cleaned_row = [clean_str(cell) for cell in row.tolist()]
                # Define the expected header sequence
                expected_headers = ['RANK', 'STUDENTS\' NAME', 'MIDTERM', 'FINALS', 'FINAL GRADE', 'EQUIVALENT', 'REMARKS', 'RANK']
                # Check if the first few elements match our expected headers
                # We'll look for a match starting from any position in the row
                for start_col in range(len(cleaned_row) - len(expected_headers) + 1):
                    matched = True
                    for j, exp_header in enumerate(expected_headers):
                        if not exp_header in cleaned_row[start_col + j]:
                            matched = False
                            break
                    if matched:
                        header_row = i
                        print(f"     Found potential header row at index {header_row}: {row_text[:50]}...")
                        break
                if header_row is not None:
                    break

        if header_row is None:
            print(f"   Could not find header row in the sheet.")
            return pd.DataFrame()

        # Read with the identified header
        df_students = pd.read_excel(class_record_path, sheet_name='Summary' if 'Summary' in pd.ExcelFile(class_record_path).sheet_names else pd.ExcelFile(class_record_path).sheet_names[0], header=header_row)
        df_students.columns = [clean_str(col).upper() for col in df_students.columns]

        # Find the necessary columns using smart_find_col or direct matching
        name_col = smart_find_col(df_students, ["STUDENTS' NAME", 'STUDENT NAME', 'NAME'])
        # The FINAL GRADE column seems to be the primary grade for ranking
        grade_col = smart_find_col(df_students, ['FINAL GRADE', 'GRADE', 'FINAL'])  # Added 'FINAL' as a possible grade column name
        rank_col = smart_find_col(df_students, ['RANK', 'RANKING'])

        # If smart_find_col fails, try direct column names
        if not name_col:
            name_col = "STUDENTS' NAME" if "STUDENTS' NAME" in df_students.columns else None
        if not grade_col:
            grade_col = 'FINAL GRADE' if 'FINAL GRADE' in df_students.columns else 'GRADE'
        if not rank_col:
            rank_col = 'RANK'

        if not name_col or not grade_col:
            print(f"   Could not find required 'NAME' and 'GRADE' columns in the sheet. Found: {list(df_students.columns)}")
            return pd.DataFrame()

        print(f"     Identified columns - Name: {name_col}, Grade: {grade_col}, Rank: {rank_col}")

        # Select only the required columns
        df_top = df_students[[name_col, grade_col]].copy()
        if rank_col and rank_col in df_students.columns:
            df_top[rank_col] = df_students[rank_col]
        else:
            df_top[rank_col] = np.nan  # Add a NaN column if rank is not found

        # Rename columns for consistency
        df_top.rename(columns={name_col: "STUDENTS' NAME", grade_col: 'GRADE', rank_col: 'RANK'}, inplace=True)

        # Convert grade to numeric
        df_top['GRADE'] = pd.to_numeric(df_top['GRADE'], errors='coerce')

        # Drop rows with missing name or grade
        df_top = df_top.dropna(subset=["STUDENTS\' NAME", 'GRADE'])

        # Sort by grade descending and assign rank if not present
        df_top = df_top.sort_values(by='GRADE', ascending=False).reset_index(drop=True)
        if pd.isna(df_top['RANK']).all():  # If all ranks are NaN
             df_top['RANK'] = range(1, len(df_top) + 1)  # Assign rank based on sorted order

        # Add the section and course code
        df_top['SECTION'] = section_from_profile  # Use the section passed from the class profile
        df_top['COURSE'] = course_code_from_profile  # Assign the course code from the class profile

        print(f"   Identified {len(df_top)} students in the class record for section {section_from_profile} (course {course_code_from_profile}).")
        # Return top 5 or all if less than 5
        return df_top.head(5)

    except Exception as e:
        print(f"❌ Error processing class record file for section {section_from_profile} (course {course_code_from_profile}): {e}")
        import traceback
        traceback.print_exc()  # Print detailed error traceback for debugging
        return pd.DataFrame()


# Helper: Normalize program names to main program codes (e.g., BSIT/BSCS)
def extract_main_program(prog_str):
    prog_str = str(prog_str).upper()
    if 'BSIT' in prog_str:
        return 'BSIT'
    elif 'BSCS' in prog_str:
        return 'BSCS'
    else:
        return prog_str  # Return as is if no known program found

# ------------------------------------------------------------
# Load LSPU Data and Extract Insights (including employment data, job paths, and student suggestions)
# ------------------------------------------------------------
def load_lspu_insights():
    """Load the LSPU data and extract common weaknesses, action_taken, recommendations, employment data, job paths, and student suggestions."""
    try:
        # Load the specific sheets we are interested in
        # Subject_Survey sheet for improvement suggestions and other feedback
        df_survey = pd.read_excel(LSPU_DATA_PATH, sheet_name='Subject_Survey')
        df_survey.columns = df_survey.columns.str.strip()
        improvement_suggestions = df_survey['subject_improvement_suggestions'].dropna().astype(str).tolist()
        # Get positive feedback as well (optional, for balance)
        positive_feedback = df_survey['subject_positive_feedback'].dropna().astype(str).tolist()
        # Academic_Deficiencies sheet for common reasons
        df_def = pd.read_excel(LSPU_DATA_PATH, sheet_name='Academic_Deficiencies')
        df_def.columns = df_def.columns.str.strip()
        deficiency_reasons = df_def['reason_of_defeciency'].dropna().astype(str).tolist()
        # Alumni_Survey and Graduate_Tracer sheets for employment data and job titles
        df_alumni = pd.read_excel(LSPU_DATA_PATH, sheet_name='Alumni_Survey')
        df_alumni.columns = df_alumni.columns.str.strip()
        # Try both possible sheet names for Graduate Tracer
        try:
            df_tracer = pd.read_excel(LSPU_DATA_PATH, sheet_name='Graduate_Tracer')
        except ValueError:
            try:
                df_tracer = pd.read_excel(LSPU_DATA_PATH, sheet_name='Graduate_Tracer_Data')
            except ValueError:
                df_tracer = pd.DataFrame()  # Create empty dataframe if neither sheet exists
        if not df_tracer.empty:
            df_tracer.columns = df_tracer.columns.str.strip()
        else:
            print("⚠️ 'Graduate_Tracer' sheet not found, using only Alumni_Survey data")

        # Combine employment data
        all_employment_data = pd.concat([df_alumni[['program', 'employment_status']],
                                        df_tracer[['program', 'employment_status']]],
                                        ignore_index=True) if not df_tracer.empty else df_alumni[['program', 'employment_status']]

        # Define employed status categories
        employed_status_values = ['Employed full-time', 'Employed part-time', 'Self-employed', 'Working full-time', 'Working part-time', 'Self-employed', 'STARTUP']
        all_employment_data['is_employed'] = all_employment_data['employment_status'].isin(employed_status_values).astype(int)

        # Calculate employment probability per program
        all_employment_data['main_program'] = all_employment_data['program'].apply(extract_main_program)
        employment_prob = all_employment_data.groupby('main_program')['is_employed'].agg(['mean', 'count']).reset_index()
        employment_prob.columns = ['main_program', 'employment_probability', 'sample_size']
        print("Employment Probabilities by Main Program:")
        print(employment_prob)

        # Extract common job titles per program (from both Alumni Survey and Graduate Tracer)
        # Filter for employed alumni from both sources
        employed_alumni = df_alumni[df_alumni['employment_status'].isin(employed_status_values)]
        if not df_tracer.empty:
            employed_tracer = df_tracer[df_tracer['employment_status'].isin(employed_status_values)]
        else:
            employed_tracer = pd.DataFrame()

        job_paths_by_program = {}
        for prog in ['BSIT', 'BSCS']:
            # Get job titles for this program from Alumni Survey (if column exists)
            alumni_jobs_for_prog = employed_alumni[employed_alumni['program'].str.contains(prog, case=False, na=False)]['job_title'].dropna().astype(str).tolist() if 'job_title' in employed_alumni.columns else []
            # Get job titles for this program from Graduate Tracer (if column exists and dataframe is not empty)
            tracer_jobs_for_prog = []
            if not employed_tracer.empty and 'job_title' in employed_tracer.columns:
                tracer_jobs_for_prog = employed_tracer[employed_tracer['program'].str.contains(prog, case=False, na=False)]['job_title'].dropna().astype(str).tolist()

            # Combine
            prog_all_jobs = alumni_jobs_for_prog + tracer_jobs_for_prog
            # Clean
            cleaned_prog_jobs = [job.strip().title() for job in prog_all_jobs if job and job != 'nan' and job.lower() != 'n/a' and job.lower() != 'na']
            # Count occurrences
            job_counts = Counter(cleaned_prog_jobs)
            # Get top 5 most common jobs
            top_jobs = [job for job, count in job_counts.most_common(5)]
            job_paths_by_program[prog] = top_jobs

        print("Top Job Paths by Main Program:")
        for prog, jobs in job_paths_by_program.items():
            print(f"{prog}: {jobs}")

        # NEW: Extract Student Suggestions per Course Code
        # Group improvement suggestions by course code
        student_suggestions_by_course = {}
        if 'course_code' in df_survey.columns and 'subject_improvement_suggestions' in df_survey.columns:
            # Clean course codes for matching
            df_survey['course_code_clean'] = df_survey['course_code'].apply(lambda x: str(x).strip().upper())
            # Group suggestions
            grouped_suggestions = df_survey.groupby('course_code_clean')['subject_improvement_suggestions'].apply(
                lambda x: [s for s in x.dropna().astype(str).tolist() if s and s != 'nan']
            ).to_dict()
            # Clean suggestions
            student_suggestions_by_course = {code: [s.strip() for s in suggestions if s and s != 'nan'] for code, suggestions in grouped_suggestions.items()}

        # NEW: Extract Subject Accomplishments (Weaknesses, Action Taken, Recommendations) per Course Code
        try:
            df_accomp = pd.read_excel(LSPU_DATA_PATH, sheet_name='Subject_Accomplishments')
            df_accomp.columns = df_accomp.columns.str.strip()
            # Clean course codes for matching
            df_accomp['course_code_clean'] = df_accomp['course_code'].apply(lambda x: str(x).strip().upper())
            # Group weaknesses, actions, and recommendations by course code
            weaknesses_by_course = df_accomp.groupby('course_code_clean')['weakness'].apply(
                lambda x: [w for w in x.dropna().astype(str).tolist() if w and w != 'nan']
            ).to_dict()
            actions_by_course = df_accomp.groupby('course_code_clean')['action_taken'].apply(
                lambda x: [a for a in x.dropna().astype(str).tolist() if a and a != 'nan']
            ).to_dict()
            accomp_recommendations_by_course = df_accomp.groupby('course_code_clean')['recommendation'].apply(
                lambda x: [r for r in x.dropna().astype(str).tolist() if r and r != 'nan']
            ).to_dict()

            # Clean lists
            weaknesses_by_course = {code: [w.strip() for w in weaknesses if w and w != 'nan'] for code, weaknesses in weaknesses_by_course.items()}
            actions_by_course = {code: [a.strip() for a in actions if a and a != 'nan'] for code, actions in actions_by_course.items()}
            accomp_recommendations_by_course = {code: [r.strip() for r in recs if r and r != 'nan'] for code, recs in accomp_recommendations_by_course.items()}
        except ValueError:
            print("⚠️ 'Subject_Accomplishments' sheet not found. Skipping related insights.")
            weaknesses_by_course = {}
            actions_by_course = {}
            accomp_recommendations_by_course = {}

        # Combine and return as a dictionary
        return {
            "improvement_suggestions": [s for s in improvement_suggestions if s and s != 'nan'],
            "positive_feedback": [pf for pf in positive_feedback if pf and pf != 'nan'],
            "deficiency_reasons": [r for r in deficiency_reasons if r and r != 'nan'],
            "employment_probabilities": employment_prob,
            "job_paths_by_program": job_paths_by_program,
            "student_suggestions_by_course": student_suggestions_by_course, # Add student suggestions
            "weaknesses_by_course": weaknesses_by_course, # Add weaknesses from Subject Accomplishments
            "actions_by_course": actions_by_course, # Add actions taken
            "accomp_recommendations_by_course": accomp_recommendations_by_course # Add recommendations from Subject Accomplishments
        }
    except Exception as e:
        print(f"Warning: Could not load LSPU insights: {e}. Using default recommendations.")
        return {
            "improvement_suggestions": [],
            "positive_feedback": [],
            "deficiency_reasons": [],
            "employment_probabilities": pd.DataFrame(columns=['main_program', 'employment_probability', 'sample_size']),
            "job_paths_by_program": {},
            "student_suggestions_by_course": {},
            "weaknesses_by_course": {},
            "actions_by_course": {},
            "accomp_recommendations_by_course": {}
        }

LSPU_INSIGHTS = load_lspu_insights()

# ------------------------------------------------------------
# Enhanced recommendation generator with student analytics and student suggestions
# ------------------------------------------------------------
def generate_recommendation(row, course_name="", program_name=""):
    course = str(course_name).upper()
    program = str(program_name).upper() # Get the program name
    pass_rate = row['PASS_RATE']
    num_def = row.get('NUM_DEF', 0)
    enrolled = row['ENROLLED']
    passed = row['PASSED']
    failed = enrolled - passed

    # Determine subject category based on course name
    subject_category = "general"
    if any(kw in course for kw in ['MOBILE', 'COMPUTING', 'LAB', 'PRACTICUM', 'INTEGRATION']):
        subject_category = "lab"
    elif any(kw in course for kw in ['THEORY', 'PRINCIPLES', 'CONCEPTS', 'WORLD']):
        subject_category = "theory"
    elif any(kw in course for kw in ['SOFTWARE', 'ENGINEERING', 'ALGORITHM', 'COMPLEXITY']):
        subject_category = "software"
    elif any(kw in course for kw in ['APPLICATION', 'DEVELOPMENT', 'MULTIMEDIA']):
        subject_category = "development"
    elif any(kw in course for kw in ['DATABASE', 'WEB', 'CLOUD']):
        subject_category = "web_dev"
    elif any(kw in course for kw in ['ALGORITHM', 'DATA STRUCTURE']):
        subject_category = "algorithm"
    elif any(kw in course for kw in ['ETHICS', 'ASSURANCE', 'SECURITY']):
        subject_category = "ethics_security"
    elif any(kw in course for kw in ['COMMUNICATION', 'ART', 'APPRECIATION']):
        subject_category = "humanities"

    # Performance level assessment
    if pass_rate >= 95:
        performance_level = "excellent"
    elif pass_rate >= 85:
        performance_level = "good"
    elif pass_rate >= 75:
        performance_level = "fair"
    else:
        performance_level = "poor"

    # Generate context-aware recommendations
    recommendations = []

    # Performance-based recommendations
    if performance_level == "excellent":
        recommendations.append(f"🎉 Excellent performance with {pass_rate}% pass rate. Continue current teaching approach.")
        if subject_category == "lab":
            recommendations.append("🧪 Maintain hands-on lab activities that contributed to success.")
        elif subject_category == "development":
            recommendations.append("💻 Excellent development skills demonstrated. Consider advanced projects.")
        elif subject_category == "software":
            recommendations.append("🛠️ Strong software engineering practices observed.")
        elif subject_category == "algorithm":
            recommendations.append("🧠 Excellent problem-solving and algorithmic thinking skills.")
    elif performance_level == "good":
        recommendations.append(f"👍 Good performance with {pass_rate}% pass rate. Minor improvements recommended.")
        if failed > 0:
            recommendations.append(f"📝 {failed} student(s) need additional support. Consider peer tutoring.")
    elif performance_level == "fair":
        recommendations.append(f"⚠️ Fair performance with {pass_rate}% pass rate. Targeted interventions needed.")
        if failed > 0:
            recommendations.append(f"👥 Schedule remedial sessions for {failed} struggling students.")
    else:  # poor performance
        recommendations.append(f"🚨 Poor performance with {pass_rate}% pass rate. Immediate action required.")
        if failed > 0:
            recommendations.append(f"🚨 Schedule mandatory consultation for {failed} students.")

    # Deficiency-based recommendations
    if num_def > 0:
        recommendations.append(f"📋 Address {int(num_def)} academic deficiency(ies) from previous semester.")
        if num_def <= 2:
            recommendations.append("📝 Individual counseling for students with deficiencies.")
        elif num_def <= 5:
            recommendations.append("👥 Form study groups for students with similar deficiencies.")
        else:
            recommendations.append("🎯 Implement comprehensive academic support program.")

    # Class size considerations
    if enrolled > 30:
        recommendations.append("👥 Large class size detected. Consider peer mentoring or teaching assistants.")
    elif enrolled < 15:
        recommendations.append("👥 Small class size allows for personalized attention.")

    # --- Employment Probability Integration ---
    # Get employment probability for the program from LSPU data
    emp_prob_df = LSPU_INSIGHTS.get("employment_probabilities", pd.DataFrame())
    if not emp_prob_df.empty and program:
        # Use the global extract_main_program function
        main_program = extract_main_program(program)
        if main_program:
            program_match = emp_prob_df[emp_prob_df['main_program'] == main_program]
            if not program_match.empty:
                emp_prob = program_match.iloc[0]['employment_probability']
                sample_size = program_match.iloc[0]['sample_size']
                if emp_prob >= 0.85:
                    recommendations.append(f"📈 Strong employment outlook ({emp_prob:.1%} based on {sample_size} alumni). Emphasize career-relevant skills.")
                elif emp_prob >= 0.70:
                    recommendations.append(f"📈 Good employment outlook ({emp_prob:.1%} based on {sample_size} alumni). Focus on practical skills.")
                else:
                    recommendations.append(f"📈 Employment outlook ({emp_prob:.1%} based on {sample_size} alumni). Consider industry partnerships and skill updates.")
            else:
                recommendations.append(f"⚠️ Employment data for main program '{main_program}' not found in LSPU records.")
        else:
            recommendations.append(f"⚠️ Could not determine main program from '{program}'. Employment data not available.")

    # --- Job Path Integration ---
    job_paths_by_program = LSPU_INSIGHTS.get("job_paths_by_program", {})
    if program:
        main_program = extract_main_program(program) # Reuse the helper function
        if main_program and main_program in job_paths_by_program:
            top_jobs = job_paths_by_program[main_program]
            if top_jobs:
                job_list_str = ", ".join(top_jobs[:3]) # Show top 3 jobs
                recommendations.append(f"💼 Potential career paths for {main_program}: {job_list_str}. Align course content with industry needs.")
            else:
                recommendations.append(f"⚠️ No specific job path data found for program '{main_program}'.")
        else:
            recommendations.append(f"⚠️ Could not determine job paths for program '{main_program}'. Data not available.")

    # --- LSPU Data Integration (Improvement Suggestions & Deficiencies) ---
    # Find relevant LSPU insights based on subject category or keywords
    relevant_insights = []

    # Get LSPU data lists
    lspu_improvements = LSPU_INSIGHTS.get("improvement_suggestions", [])
    lspu_deficiencies = LSPU_INSIGHTS.get("deficiency_reasons", [])

    # --- Match improvement suggestions ---
    category_keywords = {
        "lab": ["laboratory", "lab", "resources", "practical", "hands-on", "activities", "project"],
        "development": ["programming", "syntax", "examples", "resources", "projects", "code"],
        "web_dev": ["web", "database", "integration", "cloud", "resources", "design"],
        "algorithm": ["knowledge", "examples", "syntax", "resources", "practice"],
        "ethics_security": ["security", "assurance", "ethics", "validation", "case"],
        "theory": ["interactive", "discussion", "engagement", "online", "examples"],
        "humanities": ["communication", "art", "appreciation"],
        "general": []
    }

    keywords = category_keywords.get(subject_category, [])
    for kw in keywords:
        for s in lspu_improvements:
            if kw.lower() in s.lower() and len(s) > 10: # Avoid very short suggestions
                relevant_insights.append(f"💡 LSPU Suggestion: {s[:70]}...")
                break # Add only one per keyword to avoid clutter

    # --- Match deficiency reasons ---
    if performance_level in ["fair", "poor"]:
        for reason in lspu_deficiencies:
            if "attendance" in reason.lower():
                recommendations.append("⚠️ LSPU Data: Poor attendance is a common reason for failure. Consider attendance monitoring.")
                break
            elif "quiz" in reason.lower() or "examination" in reason.lower():
                recommendations.append("⚠️ LSPU Data: Low performance on assessments is common. Consider formative evaluations.")
                break
            elif "project" in reason.lower() or "activity" in reason.lower():
                recommendations.append("⚠️ LSPU Data: Incomplete projects/activities are common. Monitor progress closely.")
                break
            elif "no submitted" in reason.lower() or "no final" in reason.lower():
                recommendations.append("⚠️ LSPU Data: Non-submission is common. Implement reminders and check-ins.")
                break

    # Add relevant LSPU improvement suggestions to recommendations (limit to 2)
    if relevant_insights:
        recommendations.extend(random.sample(relevant_insights, min(2, len(relevant_insights))))

    # --- NEW: Student Suggestions Integration based on Course Code ---
    student_suggestions = LSPU_INSIGHTS.get("student_suggestions_by_course", {})
    # Match course code (assuming course_name is the course code like 'ITEC 106')
    course_code_clean = course.strip() # Assuming the course_name passed is the code
    if course_code_clean in student_suggestions:
        course_student_suggestions = student_suggestions[course_code_clean]
        if course_student_suggestions:
            # Add 1-2 random student suggestions for this specific course
            selected_student_suggestions = random.sample(course_student_suggestions, min(2, len(course_student_suggestions)))
            for suggestion in selected_student_suggestions:
                recommendations.append(f"🎓 Student Suggestion: {suggestion}")

    # --- NEW: Subject Accomplishments Integration (Weaknesses, Actions, Recommendations) ---
    weaknesses = LSPU_INSIGHTS.get("weaknesses_by_course", {})
    actions = LSPU_INSIGHTS.get("actions_by_course", {})
    accomp_recs = LSPU_INSIGHTS.get("accomp_recommendations_by_course", {})

    # Match course code
    if course_code_clean in weaknesses or course_code_clean in actions or course_code_clean in accomp_recs:
        # Add weaknesses if any
        if course_code_clean in weaknesses:
            course_weaknesses = weaknesses[course_code_clean]
            if course_weaknesses:
                # Add 1 weakness for context
                recommendations.append(f"🔍 Internal Review Weakness: {random.choice(course_weaknesses)}")
        # Add actions taken if any
        if course_code_clean in actions:
            course_actions = actions[course_code_clean]
            if course_actions:
                # Add 1 action taken
                recommendations.append(f"✅ Internal Review Action: {random.choice(course_actions)}")
        # Add recommendations from Subject Accomplishments if any
        if course_code_clean in accomp_recs:
            course_accomp_recs = accomp_recs[course_code_clean]
            if course_accomp_recs:
                # Add 1-2 recommendations from internal review
                selected_accomp_recs = random.sample(course_accomp_recs, min(2, len(course_accomp_recs)))
                for rec in selected_accomp_recs:
                    recommendations.append(f"📋 Internal Review Recommendation: {rec}")

    # Subject-specific dynamic recommendations (fallback if no LSPU data matched well)
    subject_recommendations = {
        "lab": [
            "🧪 Increase hands-on practice time",
            "🔧 Provide more lab equipment access",
            "🔬 Implement peer lab partners for complex experiments",
            "⚙️ Use simulation tools for complex scenarios",
            "📋 Ensure students have access to required software/applications"
        ],
        "development": [
            "💻 Set up coding review sessions",
            "👨‍💻 Encourage collaborative coding projects",
            "🔍 Implement code review practices",
            "📱 Provide access to development tools and platforms",
            "📝 Provide more examples to clarify syntax and concepts"
        ],
        "software": [
            "🛠️ Focus on project management skills",
            "📊 Implement agile development practices",
            "🤝 Emphasize team collaboration techniques",
            "🔍 Use real-world case studies"
        ],
        "web_dev": [
            "🌐 Update with latest web technologies",
            "📱 Focus on responsive design principles",
            "🔒 Emphasize security best practices",
            "🔍 Use current frameworks and tools"
        ],
        "algorithm": [
            "🧠 Practice with coding challenges",
            "📊 Visualize algorithm processes",
            "🎯 Focus on optimization techniques",
            "🔗 Connect algorithms to real applications"
        ],
        "ethics_security": [
            "🔒 Emphasize real-world security scenarios",
            "🌐 Connect concepts to current cyber threats",
            "📚 Provide practical examples of ethical dilemmas",
            "🔍 Use case studies from recent incidents"
        ],
        "humanities": [
            "💬 Encourage class discussions and debates",
            "✍️ Focus on writing and communication skills",
            "📚 Provide diverse reading materials",
            "🎯 Clarify expectations for essays and projects"
        ],
        "theory": [
            "📚 Use interactive examples",
            "🧠 Connect concepts to practical applications",
            "📝 Provide summary sheets for complex topics",
            "🔍 Use concept mapping techniques",
            "💬 Increase classroom interaction and discussion"
        ],
        "general": [
            "💬 Increase classroom interaction through Q&A sessions",
            "🎯 Set clear learning objectives for each session",
            "📅 Establish regular feedback mechanisms",
            "📚 Provide supplementary learning materials",
            "👥 Foster collaborative learning environments"
        ]
    }

    # Add 1-2 subject-specific recommendations (only if not already added via LSPU, Student, or Accomplishments)
    if not any("LSPU" in rec or "Student Suggestion" in rec or "Internal Review" in rec for rec in recommendations):
        selected_recs = random.sample(subject_recommendations.get(subject_category, subject_recommendations["general"]), min(2, len(subject_recommendations.get(subject_category, subject_recommendations["general"]))))
        recommendations.extend(selected_recs)

    # Combine all recommendations
    return " | ".join(recommendations)


# ------------------------------------------------------------
# Enhanced analytics and visualization functions with accurate gender analysis
# ------------------------------------------------------------
def generate_student_analytics(df, faculty_name, class_profile_path=None, def_report_path=None, class_record_path_map=None, academic_year=None): # Changed class_record_path to class_record_path_map, removed outstanding_path
    """Generate comprehensive student analytics and visualizations."""
    print("="*80)
    print("📊 STUDENT PERFORMANCE ANALYTICS")
    if academic_year:
        print(f"📅 ACADEMIC YEAR: {academic_year}")
    print("="*80)

    # Calculate failed students for each subject
    df['FAILED'] = df['ENROLLED'] - df['PASSED']

    # 1. Subject Strengths and Weaknesses Analysis
    print("-" * 80)
    print("🔍 SUBJECT STRENGTHS AND WEAKNESSES:")
    print("-" * 80)
    # Sort subjects by pass rate to identify strengths and weaknesses
    sorted_subjects = df.sort_values(by='PASS_RATE', ascending=True)
    weakest_subjects = sorted_subjects.head(3)
    strongest_subjects = sorted_subjects.tail(3)

    print("🔴 WEAK SUBJECTS (Lowest Pass Rates):")
    for _, row in weakest_subjects.iterrows():
        print(f"• {row['COURSE']}: {row['PASS_RATE']}% pass rate ({row['FAILED']} failed out of {row['ENROLLED']})")

    print("🟢 STRONG SUBJECTS (Highest Pass Rates):")
    for _, row in strongest_subjects.iterrows():
        print(f"• {row['COURSE']}: {row['PASS_RATE']}% pass rate ({row['PASSED']} passed out of {row['ENROLLED']})")

    # 2. Student Performance Identification
    print("-" * 80)
    print("📈 STUDENT PERFORMANCE ANALYSIS:")
    print("-" * 80)
    # Calculate individual student performance if available
    # For this we'll use a simplified approach since we don't have individual student data
    # We'll calculate based on class averages
    df['PASS_RATE_DECIMAL'] = df['PASS_RATE'] / 100
    # Identify struggling subjects
    struggling_subjects = df[df['PASS_RATE'] < 75]
    if not struggling_subjects.empty:
        print("⚠️ SUBJECTS WHERE STUDENTS ARE STRUGGLING:")
        for _, row in struggling_subjects.iterrows():
            print(f"• {row['COURSE']}: {row['PASS_RATE']}% pass rate")
    else:
        print("✅ No subjects are struggling (pass rate >= 75%).")

    # 3. Gender Analysis (Now extracting from Class Academic Profile with accurate column mapping)
    print("-" * 80)
    print("👥 GENDER ANALYSIS:")
    print("-" * 80)
    if class_profile_path:
        try:
            # Read the original class profile to get the gender data
            df_raw = pd.read_excel(class_profile_path, header=None)
            # Look for the actual header row (usually contains terms like "PROGRAM", "COURSE", "ENROLLED")
            header_row = None
            for i in range(min(10, df_raw.shape[0])):
                row = df_raw.iloc[i].astype(str).str.upper()
                row_text = ' '.join(row.values).upper()
                if any(word in row_text for word in ['PROGRAM', 'COURSE', 'ENROLLED', 'PASSED', 'STUDENTS']):
                    header_row = i
                    break

            if header_row is None:
                header_row = 4

            # Read with proper header
            df_gender = pd.read_excel(class_profile_path, header=header_row)
            # Remove rows that are clearly not data rows (like "TOTAL", "PERCENTAGE", "Prepared by")
            df_gender = df_gender[~df_gender.iloc[:, 0].astype(str).str.upper().isin(['TOTAL', 'PERCENTAGE', 'PREPARED BY:', 'CHECKED:', 'APPROVED BY:', 'NOTED:'])]

            # ACCURATE GENDER COLUMN DETECTION BASED ON ACTUAL FILE STRUCTURE
            # Based on the image you shared, the columns are:
            # 0: PROGRAM, YEAR, AND SECTION
            # 1: COURSE
            # 2: STUDENTS ENROLLED IN CLASS (Male)
            # 3: STUDENTS ENROLLED IN CLASS (Female)
            # 4: STUDENTS ENROLLED IN CLASS (Total)
            # 5: COMPUTER-STUDENT RATIO
            # 6: PASSED STUDENTS (Male)
            # 7: PASSED STUDENTS (Female)
            # 8: PASSED STUDENTS (Total)
            # 9: FAILED STUDENTS (Male)
            # 10: FAILED STUDENTS (Female)
            # 11: FAILED STUDENTS (Total)
            # Check if we have enough columns for accurate gender analysis
            if df_gender.shape[1] >= 8:  # Need at least 8 columns
                # Calculate totals from the actual column positions
                # Use pd.to_numeric to ensure data types are correct before summing and handle non-numeric gracefully
                total_male_enrolled_raw = df_gender.iloc[:, 2]
                total_female_enrolled_raw = df_gender.iloc[:, 3]
                total_male_passed_raw = df_gender.iloc[:, 6]
                total_female_passed_raw = df_gender.iloc[:, 7]

                # Convert to numeric, setting errors to NaN
                total_male_enrolled_numeric = pd.to_numeric(total_male_enrolled_raw, errors='coerce')
                total_female_enrolled_numeric = pd.to_numeric(total_female_enrolled_raw, errors='coerce')
                total_male_passed_numeric = pd.to_numeric(total_male_passed_raw, errors='coerce')
                total_female_passed_numeric = pd.to_numeric(total_female_passed_raw, errors='coerce')

                # Sum, ignoring NaN values
                total_male_enrolled = total_male_enrolled_numeric.sum()
                total_female_enrolled = total_female_enrolled_numeric.sum()
                total_male_passed = total_male_passed_numeric.sum()
                total_female_passed = total_female_passed_numeric.sum()

                # Calculate total students (should equal male + female)
                total_students = total_male_enrolled + total_female_enrolled

                print("📊 Gender Distribution:")
                print(f"• Male Students: {int(total_male_enrolled) if pd.notna(total_male_enrolled) else 0}")
                print(f"• Female Students: {int(total_female_enrolled) if pd.notna(total_female_enrolled) else 0}")
                print(f"• Total Students: {int(total_students) if pd.notna(total_students) else 0}")

                # Validate that our calculation matches the total in the file
                expected_total_raw = df_gender.iloc[:, 4]
                expected_total_numeric = pd.to_numeric(expected_total_raw, errors='coerce')
                expected_total = expected_total_numeric.sum()
                if pd.notna(expected_total) and expected_total > 0:
                    print(f"• Expected Total (from file): {int(expected_total)}")
                    if abs(total_students - expected_total) > 5:  # Allow small discrepancy
                        print(f"⚠️ Warning: Calculated total ({total_students}) differs significantly from file total ({expected_total})")

                # Calculate pass rates by gender
                # Handle potential division by zero or NaN
                male_pass_rate = (total_male_passed / total_male_enrolled * 100) if total_male_enrolled > 0 and pd.notna(total_male_passed) and pd.notna(total_male_enrolled) else 0
                female_pass_rate = (total_female_passed / total_female_enrolled * 100) if total_female_enrolled > 0 and pd.notna(total_female_passed) and pd.notna(total_female_enrolled) else 0

                print("📈 Pass Rates by Gender:")
                print(f"• Male: {male_pass_rate:.2f}% ({int(total_male_passed) if pd.notna(total_male_passed) else 0} passed out of {int(total_male_enrolled) if pd.notna(total_male_enrolled) else 0})")
                print(f"• Female: {female_pass_rate:.2f}% ({int(total_female_passed) if pd.notna(total_female_passed) else 0} passed out of {int(total_female_enrolled) if pd.notna(total_female_enrolled) else 0})")

                if male_pass_rate > female_pass_rate:
                    print(f"🏆 Males are performing better with {male_pass_rate:.2f}% vs {female_pass_rate:.2f}%")
                elif female_pass_rate > male_pass_rate:
                    print(f"🏆 Females are performing better with {female_pass_rate:.2f}% vs {male_pass_rate:.2f}%")
                else:
                    print("✅ Both genders are performing equally well")

                # Top performing sections by gender
                print("🏆 TOP PERFORMING SECTIONS BY GENDER:")
                # Calculate pass rates for each section using numeric data
                if df_gender.shape[1] > 6 and df_gender.shape[1] > 2:  # Check if columns exist
                    male_passed_numeric = pd.to_numeric(df_gender.iloc[:, 6], errors='coerce')
                    male_enrolled_numeric = pd.to_numeric(df_gender.iloc[:, 2], errors='coerce')
                    # Calculate pass rate: (passed / enrolled) * 100, handling division by zero and NaN
                    male_pass_rates_raw = (male_passed_numeric / male_enrolled_numeric) * 100
                    male_pass_rates = male_pass_rates_raw.fillna(0) # Replace NaN with 0 for sections with no enrolled students
                else:
                    male_pass_rates = pd.Series([0] * len(df_gender))

                if df_gender.shape[1] > 7 and df_gender.shape[1] > 3:  # Check if columns exist
                    female_passed_numeric = pd.to_numeric(df_gender.iloc[:, 7], errors='coerce')
                    female_enrolled_numeric = pd.to_numeric(df_gender.iloc[:, 3], errors='coerce')
                    female_pass_rates_raw = (female_passed_numeric / female_enrolled_numeric) * 100
                    female_pass_rates = female_pass_rates_raw.fillna(0) # Replace NaN with 0
                else:
                    female_pass_rates = pd.Series([0] * len(df_gender))

                # Add section and course info for display
                sections = df_gender.iloc[:, 0] if df_gender.shape[1] > 0 else pd.Series(["Unknown"] * len(df_gender))
                courses = df_gender.iloc[:, 1] if df_gender.shape[1] > 1 else pd.Series(["Unknown"] * len(df_gender))

                # Create a temporary dataframe for analysis
                temp_df = pd.DataFrame({
                    'SECTION': sections,
                    'COURSE': courses,
                    'MALE_PASS_RATE': male_pass_rates,
                    'FEMALE_PASS_RATE': female_pass_rates
                })

                # Find top sections for males (excluding 0% pass rate if all are 0, or handle gracefully)
                top_male_sections = temp_df.nlargest(3, 'MALE_PASS_RATE')
                print("🥇 Top Performing Sections for Males:")
                for _, row in top_male_sections.iterrows():
                    if row['MALE_PASS_RATE'] > 0: # Only show if pass rate is greater than 0
                        print(f"• {row['SECTION']} - {row['COURSE']}: {row['MALE_PASS_RATE']:.2f}% pass rate")
                    else:
                        print(f"• No male sections with pass rates > 0% found.")
                        break # Exit loop after printing message if no valid sections

                # Find top sections for females
                top_female_sections = temp_df.nlargest(3, 'FEMALE_PASS_RATE')
                print("🥇 Top Performing Sections for Females:")
                for _, row in top_female_sections.iterrows():
                    if row['FEMALE_PASS_RATE'] > 0: # Only show if pass rate is greater than 0
                        print(f"• {row['SECTION']} - {row['COURSE']}: {row['FEMALE_PASS_RATE']:.2f}% pass rate")
                    else:
                        print(f"• No female sections with pass rates > 0% found.")
                        break # Exit loop after printing message if no valid sections

            else:
                print("⚠️ Not enough columns found for accurate gender analysis in the Class Academic Profile.")
                print("   Expected at least 8 columns with gender-specific data.")
        except Exception as e:
            print(f"⚠️ Could not extract gender  {e}")
            print("   Detailed Error Info: ", str(e))
            print("ℹ️ Gender analysis requires properly formatted class profile with numeric gender columns.")
    else:
        print("ℹ️ Note: Gender analysis requires the Class Academic Profile file path.")
        print("   If available, upload student records with gender information to enable this analysis.")

    # 4. Student Status Overview
    print("-" * 80)
    print("📋 STUDENT STATUS OVERVIEW:")
    print("-" * 80)
    total_enrolled = df['ENROLLED'].sum()
    total_passed = df['PASSED'].sum()
    total_failed = total_enrolled - total_passed
    total_deficiencies = df['NUM_DEF'].sum()

    print(f"• Total Students Enrolled: {total_enrolled}")
    print(f"• Total Students Passed: {total_passed}")
    print(f"• Total Students Failed: {total_failed}")
    print(f"• Total Academic Deficiencies: {total_deficiencies}")

    # 5. Alarming Students Identification
    print("-" * 80)
    print("🚨 ALARMING SUBJECTS IDENTIFICATION:")
    print("-" * 80)
    # Identify subjects with high failure rates and high deficiencies
    alarming_subjects = df[
        (df['PASS_RATE'] < 70) |
        (df['NUM_DEF'] > df['ENROLLED'] * 0.1)  # More than 10% have deficiencies
    ]
    if not alarming_subjects.empty:
        print("⚠️ ALARMING SUBJECTS REQUIRING IMMEDIATE ATTENTION:")
        for _, row in alarming_subjects.iterrows():
            issues = []
            if row['PASS_RATE'] < 70:
                issues.append(f"Low pass rate: {row['PASS_RATE']}%")
            if row['NUM_DEF'] > row['ENROLLED'] * 0.1:
                issues.append(f"High deficiencies: {row['NUM_DEF']}")
            print(f"• {row['COURSE']}: {', '.join(issues)}")
    else:
        print("✅ No alarming subjects identified at this time.")

    # 6. Summary of Struggling Subjects
    print("-" * 80)
    print("📋 SUMMARY: MOST STRUGGLED SUBJECTS:")
    print("-" * 80)
    struggling_summary = df[df['PASS_RATE'] < 80].sort_values(by='PASS_RATE', ascending=True)
    if not struggling_summary.empty:
        for _, row in struggling_summary.iterrows():
            print(f"• {row['COURSE']}: {row['PASS_RATE']}% pass rate")
    else:
        print("✅ All subjects are performing above 80% pass rate.")

    # 7. Specific Student Identification (Alarming and Outstanding)
    print("-" * 80)
    print("👤 SPECIFIC STUDENT IDENTIFICATION:")
    print("-" * 80)

    # --- NEW SECTION: OUTSTANDING STUDENTS (Top Performers) - From Class Record ---
    print("🌟 OUTSTANDING STUDENTS (Top Performers - From Class Record):")
    if class_record_path_map: # Use the map instead of a single path
        global CLASS_RECORD_STUDENTS
        # Iterate through the processed DataFrame (df) which contains course and section info
        for index, row in df.iterrows():
            course_code_from_profile = row['COURSE']
            section_from_profile = row['PROGRAM'] # Assuming 'PROGRAM' column holds the section info like 'BSIT IIA'
            print(f"  Looking for course '{course_code_from_profile}' and section '{section_from_profile}' in Class Record Map.")
            # Find the corresponding class record file path based on the course code
            matching_record_path = None
            for record_path, record_course_code in class_record_path_map.items():
                if record_course_code == course_code_from_profile:
                    matching_record_path = record_path
                    break

            if matching_record_path:
                 print(f"    Found matching Class Record file: {os.path.basename(matching_record_path)} for course {course_code_from_profile}")
                 # Process the specific section within that file
                 section_students = process_class_record_students(matching_record_path, course_code_from_profile, section_from_profile)
                 if not section_students.empty:
                     print(f"    Top 5 performers for {course_code_from_profile} ({section_from_profile}):")
                     for _, s_row in section_students.iterrows():
                         rank = int(s_row['RANK'])
                         name = s_row["STUDENTS' NAME"]
                         section = s_row['SECTION']
                         grade = s_row['GRADE']
                         print(f"      - Rank {rank}: {name} ({section}) - Grade: {grade}")
                 else:
                     print(f"    No outstanding students data found from the class record for {course_code_from_profile} ({section_from_profile}) or error processing the file.")
            else:
                print(f"    No Class Record file found for course {course_code_from_profile} in the uploaded files.")
    else:
        print("  No Class Record files provided for identifying outstanding students.")

    # --- REMOVED SECTION: OUTSTANDING STUDENTS (From Separate List) ---

    if def_report_path:
        try:
            # Try to read as DOCX
            doc = Document(def_report_path)
            tables = []
            for table in doc.tables:
                data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                # Ensure table has header and rows before creating DataFrame
                if data and len(data) > 1 and len(data[0]) > 0:
                    try:
                        df_def = pd.DataFrame(data[1:], columns=data[0])
                    except Exception:
                        df_def = pd.DataFrame(data)
                    break # Process only the first table found

            if 'df_def' in locals() and df_def is not None and not df_def.empty:
                # Process deficiency report to identify specific students
                df_def.columns = [clean_str(c).upper() for c in df_def.columns]
                # Look for student name, course, and grade columns
                name_col = smart_find_col(df_def, ['NAME', 'STUDENT', 'NAME OF STUDENT'])
                course_col = smart_find_col(df_def, ['COURSE', 'SUBJECT'])
                grade_col = smart_find_col(df_def, ['GRADE', 'SCORE'])
                reason_col = smart_find_col(df_def, ['REASON', 'DEFICIENCY'])

                if name_col and course_col and grade_col:
                    # --- IMPROVED GRADE PARSING ---
                    df_def['GRADE_NUMERIC'] = pd.to_numeric(df_def[grade_col].str.extract(r'(\d+\.?\d*)')[0], errors='coerce')
                    df_def['GRADE_RAW'] = df_def[grade_col].str.upper()

                    # Identify alarming students based on grades using both numeric and raw text checks
                    df_def['STATUS'] = df_def.apply(
                        lambda row: 'FAILED' if row['GRADE_NUMERIC'] == 5.0 or '5.00' in row['GRADE_RAW'] or 'FAILED' in row['GRADE_RAW'] else
                                'INCOMPLETE' if 'INC' in row['GRADE_RAW'] else 'AT RISK',
                        axis=1
                    )

                    # Group by status to show all failed, incomplete, at-risk students
                    try:
                        df_def['VALID_NAME'] = df_def[name_col].apply(is_valid_full_name)
                    except Exception:
                        df_def['VALID_NAME'] = True
                    print("🔴 FAILED STUDENTS (5.00 or FAILED):")
                    failed_students = df_def[(df_def['STATUS'] == 'FAILED') & (df_def['VALID_NAME'])]
                    if not failed_students.empty:
                        for _, row in failed_students.iterrows():
                            print(f"  • {row[name_col]} - {row[course_col]} - {row[grade_col]}")
                    else:
                        print("  No failed students found.")

                    print("📋 INCOMPLETE STUDENTS (INC):")
                    incomplete_students = df_def[(df_def['STATUS'] == 'INCOMPLETE') & (df_def['VALID_NAME'])]
                    if not incomplete_students.empty:
                        for _, row in incomplete_students.iterrows():
                            print(f"  • {row[name_col]} - {row[course_col]} - {row[grade_col]}")
                    else:
                        print("  No incomplete students found.")

                    # Identify students with multiple issues
                    print("⚠️ STUDENTS WITH MULTIPLE ISSUES:")
                    student_counts = df_def[df_def['VALID_NAME']][name_col].value_counts()
                    multiple_issues = student_counts[student_counts > 1]
                    if not multiple_issues.empty:
                        for student, count in multiple_issues.items():
                            print(f"  • {student} - {count} subjects with issues")
                            student_records = df_def[df_def[name_col] == student]
                            for _, row in student_records.iterrows():
                                print(f"    - {row[course_col]}: {row[grade_col]} ({row['STATUS']})")
                    else:
                        print("  No students with multiple issues found.")

                    # ML-based risk assessment
                    print("🤖 AI PREDICTED HIGH-RISK STUDENTS:")
                    if reason_col:
                        high_risk = df_def[df_def[reason_col].str.contains('attendance|no quiz|no laboratories|no final|no submitted|no midterm|few attendance', case=False, na=False)]
                        high_risk = high_risk[high_risk['VALID_NAME']]
                        if not high_risk.empty:
                            for _, row in high_risk.iterrows():
                                print(f"  • {row[name_col]} - {row[course_col]} - High risk due to: {row[reason_col]}")
                        else:
                            print("  No high-risk students identified by AI.")
                    else:
                        print("  Reason column not found for AI risk assessment.")
                else:
                    print("⚠️ Could not find required columns in deficiency report for specific student identification.")
                    print("   Looking for columns like 'NAME', 'COURSE', 'GRADE'.")
            else:
                print("⚠️ Could not find a table in the deficiency report.")
                print("   Please ensure the deficiency report has proper table structure with student names, courses, and grades.")
        except Exception as e:
            print(f"⚠️ Could not process deficiency report for specific student identification: {e}")
            print("ℹ️ Please ensure the deficiency report has proper table structure with student names, courses, and grades.")
    else:
        print("ℹ️ Deficiency report not provided. Specific student identification requires the deficiency report.")

    # 8. Generate Visualizations
    print("-" * 80)
    print("📊 GENERATING VISUALIZATIONS...")
    print("-" * 80)

    # Create subplots for comprehensive visualization
    fig = make_subplots(
        rows=2, cols=2,
        subplot_titles=('Pass Rate by Subject', 'Enrollment vs Pass Rate', 'Performance Distribution', 'Deficiencies by Subject'),
        specs=[[{"type": "bar"}, {"type": "scatter"}],
              [{"type": "pie"}, {"type": "bar"}]]
    )

    # Bar chart: Pass Rate by Subject
    fig.add_trace(
        go.Bar(x=df['COURSE'], y=df['PASS_RATE'], name='Pass Rate (%)'),
        row=1, col=1
    )

    # Scatter: Enrollment vs Pass Rate
    fig.add_trace(
        go.Scatter(x=df['ENROLLED'], y=df['PASS_RATE'], mode='markers',
                  text=df['COURSE'], name='Enrollment vs Pass Rate'),
        row=1, col=2
    )

    # Pie chart: Performance Distribution
    performance_data = {
        'Passed': total_passed,
        'Failed': total_failed
    }
    fig.add_trace(
        go.Pie(labels=list(performance_data.keys()), values=list(performance_data.values()),
              name='Performance Distribution'),
        row=2, col=1
    )

    # Bar chart: Deficiencies by Subject
    fig.add_trace(
        go.Bar(x=df['COURSE'], y=df['NUM_DEF'], name='Deficiencies', marker_color='red'),
        row=2, col=2
    )

    fig.update_layout(height=800, title_text=f"Comprehensive Analytics Dashboard - {faculty_name}", showlegend=False)
    fig.show()

    # Additional visualization: Performance by subject category
    print("📈 PERFORMANCE BY SUBJECT CATEGORY:")
    df['SUBJECT_CATEGORY'] = df['COURSE'].apply(lambda x:
        'Lab' if any(kw in x.upper() for kw in ['MOBILE', 'COMPUTING', 'LAB', 'PRACTICUM', 'INTEGRATION']) else
        'Theory' if any(kw in x.upper() for kw in ['THEORY', 'PRINCIPLES', 'CONCEPTS', 'WORLD']) else
        'Software' if any(kw in x.upper() for kw in ['SOFTWARE', 'ENGINEERING', 'ALGORITHM', 'COMPLEXITY']) else
        'Development' if any(kw in x.upper() for kw in ['APPLICATION', 'DEVELOPMENT', 'MULTIMEDIA']) else
        'Web/DB' if any(kw in x.upper() for kw in ['DATABASE', 'WEB', 'CLOUD']) else
        'Algorithm' if any(kw in x.upper() for kw in ['ALGORITHM', 'DATA STRUCTURE']) else
        'Security' if any(kw in x.upper() for kw in ['ETHICS', 'ASSURANCE', 'SECURITY']) else
        'Humanities' if any(kw in x.upper() for kw in ['COMMUNICATION', 'ART', 'APPRECIATION']) else
        'General'
    )
    category_performance = df.groupby('SUBJECT_CATEGORY')['PASS_RATE'].mean().reset_index()
    fig2 = px.bar(category_performance, x='SUBJECT_CATEGORY', y='PASS_RATE',
                  title='Average Pass Rate by Subject Category',
                  labels={'PASS_RATE': 'Average Pass Rate (%)', 'SUBJECT_CATEGORY': 'Subject Category'})
    fig2.show()

    # 9. === NEW AI-GENERATED SUMMARY SECTION ===
    print("=" * 80)
    print("Summary of Findings (AI Generated):")
    print("=" * 80)

    # Calculate overall metrics for the AI summary
    total_students_enrolled = df['ENROLLED'].sum()
    total_students_passed = df['PASSED'].sum()
    total_students_failed = total_students_enrolled - total_students_passed
    total_academic_deficiencies = int(df['NUM_DEF'].sum())

    # Prevent division by zero
    average_pass_rate = (total_students_passed / total_students_enrolled * 100) if total_students_enrolled > 0 else 0
    fail_rate = (total_students_failed / total_students_enrolled * 100) if total_students_enrolled > 0 else 0

    # --- Attempt to parse deficiency report for specific "At Risk" student count ---
    at_risk_student_count = 0
    if def_report_path: # Using the def_report_path passed to the function
        try:
            # Try to read the deficiency report DOCX to count unique failing students
            doc = Document(def_report_path)
            if doc.tables:
                # Assume the first table is the one we need
                table = doc.tables[0]
                data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                if data and len(data) > 1:
                    df_def_temp = pd.DataFrame(data[1:], columns=data[0])
                    # Standardize column names for the temp DF
                    standardized_columns = {col: clean_str(col).upper() for col in df_def_temp.columns}
                    df_def_temp.rename(columns=standardized_columns, inplace=True)

                    # Find the name and grade columns
                    name_col_temp = smart_find_col(df_def_temp, ['NAME', 'STUDENT', 'NAME OF STUDENT'])
                    grade_col_temp = smart_find_col(df_def_temp, ['GRADE', 'SCORE'])

                    if name_col_temp and grade_col_temp:
                        # --- IMPROVED STUDENT COUNTING ---
                        # Filter for rows where grade indicates failure (5.00, FAILED)
                        failed_mask = df_def_temp[grade_col_temp].str.contains(r"5\.00|FAILED", case=False, na=False)
                        failed_students_df = df_def_temp[failed_mask]
                        # Count unique student names in the failed subset
                        at_risk_student_count = failed_students_df[name_col_temp].nunique()
                    # else: required columns not found, keep count at 0
                # else: empty table, keep count at 0
            # else: no tables, keep count at 0
        except Exception as e:
            # If anything goes wrong in parsing, default to calculating from aggregate failures
            # This is a fallback, less precise than counting unique students in the report
            print(f"    (Note: Could not parse deficiency report for unique at-risk count: {e}. Using aggregate failed count.)")
            at_risk_student_count = total_students_failed
    else:
        # If no def report path, fall back to aggregate failed count
        at_risk_student_count = total_students_failed


    # --- Generate the text summary ---
    # Note: The phrase "grade data extracted from the '1B' sheets of 836 students" is specific to an example.
    # We'll adapt it to reflect the data we actually have.
    # The number of subjects processed is simply the number of rows in df.
    num_subjects_processed = len(df)

    print(f"Based on the uploaded class records summarizing performance across {num_subjects_processed} subjects "
          f"and the deficiency report, a total of {total_students_enrolled} students are enrolled. "
          f"There are {at_risk_student_count} students flagged as 'At Risk' (Failed), "
          f"representing {fail_rate:.1f}% of the cohort. "
          f"The overall average pass rate is {average_pass_rate:.2f}%, and the fail rate is {fail_rate:.1f}%. "
          f"Additionally, {total_academic_deficiencies} academic deficiencies were recorded. "
          f"This summary was generated by processing the Class Academic Profile and Deficiency Report.")

    # ===============================

    # 10. Advanced Analytics and Specific Visualizations
    print("="*80)
    print("🚀 ADVANCED ANALYTICS AND SPECIFIC VISUALIZATIONS")
    print("="*80)

    # Determine the subject code from the processed DataFrame (df)
    # We'll use the first course found as the subject of interest for detailed analysis
    if not df.empty:
        first_course_code = df.iloc[0]['COURSE']
        print(f"📊 Analyzing detailed metrics for: {first_course_code}")
    else:
        print("📊 No course data found for detailed analysis.")
        first_course_code = "UNKNOWN"

    # 10.1. Program Pass Rate Comparison Bar Chart
    print("📊 1. Program Pass Rate Comparison")
    program_summary = df.groupby('PROGRAM').agg({
        'ENROLLED': 'sum',
        'PASSED': 'sum',
        'FAILED': 'sum'
    }).reset_index()
    program_summary['Pass Rate'] = (program_summary['PASSED'] / program_summary['ENROLLED'] * 100).round(2)

    # Add Overall
    overall_summary = pd.DataFrame({
        'PROGRAM': ['Overall'],
        'ENROLLED': [total_students_enrolled],
        'PASSED': [total_students_passed],
        'FAILED': [total_students_failed],
        'Pass Rate': [average_pass_rate]
    })
    program_summary = pd.concat([program_summary, overall_summary], ignore_index=True)

    fig_program = px.bar(program_summary, x='PROGRAM', y='Pass Rate',
                        title='Pass Rate Comparison by Program',
                        labels={'Pass Rate': 'Pass Rate (%)', 'PROGRAM': 'Program'},
                        color='PROGRAM', color_discrete_sequence=px.colors.qualitative.Set3)
    fig_program.update_layout(yaxis_range=[0, 100])
    fig_program.show()

    # 10.2. Course Effectiveness Clusters Scatter Plot
    print("📊 2. Course Effectiveness Clusters")
    # Prepare data for clustering (enrollment and pass rate)
    course_cluster_data = df[['ENROLLED', 'PASS_RATE']].copy()

    # K-Means Clustering
    kmeans = KMeans(n_clusters=3, random_state=42)
    course_cluster_data['Cluster'] = kmeans.fit_predict(course_cluster_data)

    # Create scatter plot
    fig_course_cluster = px.scatter(course_cluster_data, x='ENROLLED', y='PASS_RATE',
                                    color='Cluster',
                                    title='Course Effectiveness Clusters',
                                    labels={'ENROLLED': 'Total Enrollment', 'PASS_RATE': 'Pass Rate (%)'},
                                    color_discrete_sequence=px.colors.qualitative.Set2)
    fig_course_cluster.update_layout(
        legend_title_text='Cluster',
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=1.02,
            xanchor="right",
            x=1
        )
    )
    # Add cluster descriptions to legend
    cluster_descriptions = {
        0: "High Effectiveness",
        1: "Medium Effectiveness",
        2: "Low Effectiveness"
    }
    for i, desc in cluster_descriptions.items():
        fig_course_cluster.add_annotation(
            xref="paper", yref="paper",
            x=0.02, y=0.98 - i*0.03,
            text=f"Cluster {i} ({desc})",
            showarrow=False,
            font=dict(size=10),
            bgcolor="white",
            bordercolor="black",
            borderwidth=1
        )
    fig_course_cluster.show()

    # 10.3. Student Performance Clusters for the first course in the uploaded data
    print(f"📊 3. Student Performance Clusters ({first_course_code})")
    relevant_students_df = pd.DataFrame() # Initialize as empty DataFrame
    if def_report_path and first_course_code != "UNKNOWN":
        try:
            doc = Document(def_report_path)
            if doc.tables:
                table = doc.tables[0]
                data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                if data and len(data) > 1:
                    df_def_temp = pd.DataFrame(data[1:], columns=data[0])
                    # Standardize column names
                    standardized_columns = {col: clean_str(col).upper() for col in df_def_temp.columns}
                    df_def_temp.rename(columns=standardized_columns, inplace=True)

                    # Find relevant columns
                    name_col_temp = smart_find_col(df_def_temp, ['NAME', 'STUDENT'])
                    course_col_temp = smart_find_col(df_def_temp, ['COURSE', 'SUBJECT'])
                    grade_col_temp = smart_find_col(df_def_temp, ['GRADE', 'SCORE'])

                    if name_col_temp and course_col_temp and grade_col_temp:
                        # Filter for the specific course code found in the class profile
                        relevant_filter = (
                            df_def_temp[course_col_temp].str.contains(first_course_code, case=False, na=False)
                        )
                        relevant_students_df = df_def_temp[relevant_filter].copy() # Use .copy()
                        if not relevant_students_df.empty:
                            print(f"  Found {len(relevant_students_df)} students with deficiencies for {first_course_code}.")
                            # Note: The deficiency report likely contains final grades or deficiency reasons,
                            # not separate Midterm/Final grades. We'll proceed assuming a 'GRADE' column exists.
                            # Attempt to parse the grade column as a numeric 'Final Grade'
                            relevant_students_df['Final Grade'] = pd.to_numeric(relevant_students_df[grade_col_temp].str.extract(r'(\d+\.\d+)')[0], errors='coerce')

                            # Drop rows with invalid final grades (this might be problematic if grades are 'INC', 'DRP', etc.)
                            relevant_students_df = relevant_students_df.dropna(subset=['Final Grade'])
                            print(f"  After parsing grades, {len(relevant_students_df)} students have valid numeric grades.")

                            if not relevant_students_df.empty:
                                # For clustering, we need two features. Since we likely only have one grade,
                                # we'll create a dummy second feature (e.g., a constant or the grade itself again)
                                # This isn't ideal for clustering but allows the code to run.
                                # A better approach might be to use other features if available (e.g., number of deficiencies, reason).
                                relevant_students_df['Dummy Feature'] = relevant_students_df['Final Grade'] # Use grade as a second feature

                                # K-Means Clustering on the available grade data and the dummy feature
                                student_cluster_data = relevant_students_df[['Final Grade', 'Dummy Feature']].copy()
                                kmeans_student = KMeans(n_clusters=3, random_state=42)
                                student_cluster_data['Cluster'] = kmeans_student.fit_predict(student_cluster_data)

                                # Create scatter plot
                                fig_student_cluster = px.scatter(student_cluster_data, x='Final Grade', y='Dummy Feature',
                                                                color='Cluster',
                                                                title=f'Student Performance Clusters ({first_course_code})',
                                                                labels={'Final Grade': 'Grade', 'Dummy Feature': 'Grade (Copy)'},
                                                                color_discrete_sequence=px.colors.qualitative.Set2)
                                fig_student_cluster.update_layout(
                                    legend_title_text='Cluster',
                                    legend=dict(
                                        orientation="h",
                                        yanchor="bottom",
                                        y=1.02,
                                        xanchor="right",
                                        x=1
                                    )
                                )
                                # Add cluster descriptions to legend
                                cluster_descriptions_student = {
                                    0: "Cluster A",
                                    1: "Cluster B",
                                    2: "Cluster C"
                                }
                                for i, desc in cluster_descriptions_student.items():
                                    fig_student_cluster.add_annotation(
                                        xref="paper", yref="paper",
                                        x=0.02, y=0.98 - i*0.03,
                                        text=f"Cluster {i} ({desc})",
                                        showarrow=False,
                                        font=dict(size=10),
                                        bgcolor="white",
                                        bordercolor="black",
                                        borderwidth=1
                                    )
                                fig_student_cluster.show()
                            else:
                                print(f"  No students with valid numeric grades found for {first_course_code} in deficiency report.")
                        else:
                            print(f"  No students found for {first_course_code} in deficiency report.")
                    else:
                        print(f"  Required columns not found for {first_course_code} student analysis. Found: name={name_col_temp}, course={course_col_temp}, grade={grade_col_temp}")
                else:
                    print(f"  Empty table for {first_course_code} student analysis.")
            else:
                print(f"  No tables found for {first_course_code} student analysis.")
        except Exception as e:
            print(f"  Error processing deficiency report for {first_course_code} student analysis: {e}")
    else:
        print(f"  Deficiency report not provided or no course data found for {first_course_code} student analysis.")

    # 10.4. Predictive Model for Finals (This is difficult without Midterm/Final grade pairs)
    print(f"📊 4. Predictive Model for Finals ({first_course_code})")
    if not relevant_students_df.empty and 'Final Grade' in relevant_students_df.columns:
        # Since we likely don't have Midterm grades, we cannot build a predictive model for Finals.
        # We could potentially use other features (e.g., number of deficiencies, reason) to predict pass/fail,
        # but that requires more complex feature engineering.
        # For now, we'll just state that the data is insufficient.
        print(f"  Cannot generate predictive model for {first_course_code} Finals grade.")
        print(f"  Reason: Deficiency report likely contains final grade/deficiency reason, not Midterm grade required for prediction.")
    else:
        print(f"  No student data available for predictive modeling for {first_course_code}.")

    # 10.5. Alarming vs. Outstanding Students (Based on Deficiency Report Grades)
    print(f"📊 5. Alarming vs. Outstanding Students ({first_course_code})")
    if not relevant_students_df.empty and 'Final Grade' in relevant_students_df.columns:
        # Create a new column for status based on the numeric grade
        # Assuming standard LPU grading: 5.00 is failing, < 3.00 might be concerning, >= 3.00 is passing
        # Adjust thresholds as needed based on LPU standards
        def categorize_grade(grade):
            if pd.isna(grade):
                return 'Unknown'
            if grade <= 5.00 and grade >= 3.00:
                return 'Alarming'
            elif grade < 3.00:
                return 'Outstanding'
            else: # Grade > 5.00 (shouldn't happen with 5.00 max, but just in case)
                return 'Other'

        relevant_students_df['Status'] = relevant_students_df['Final Grade'].apply(categorize_grade)

        # Create scatter plot
        fig_alarm_out = go.Figure()

        # Add points for different statuses
        for status in ['Outstanding', 'Alarming', 'Other', 'Unknown']:
            subset = relevant_students_df[relevant_students_df['Status'] == status]
            if not subset.empty:
                fig_alarm_out.add_trace(go.Scatter(
                    x=subset['Final Grade'], # X-axis could be grade or another metric if available
                    y=[0] * len(subset), # Y-axis is arbitrary for this visualization
                    mode='markers',
                    name=status,
                    marker=dict(
                        color='green' if status == 'Outstanding' else ('red' if status == 'Alarming' else ('orange' if status == 'Other' else 'gray')),
                        size=10
                    ),
                    text=subset[name_col_temp] if name_col_temp else None # Hover text with student names
                ))

        fig_alarm_out.update_layout(
            title=f'Alarming vs. Outstanding Students ({first_course_code})',
            xaxis_title='Grade',
            yaxis_title='N/A',
            showlegend=True
        )
        fig_alarm_out.show()
    else:
        print(f"  No student data available for alarming vs. outstanding analysis for {first_course_code}.")

    # 10.6. Predicted Risk Distribution Histogram (Based on Deficiency Report Grades)
    print(f"📊 6. Predicted Risk Distribution ({first_course_code})")
    if not relevant_students_df.empty and 'Final Grade' in relevant_students_df.columns:
        # Create bins for final grades based on LPU standard (e.g., 1.00 to 5.00)
        # Adjust bins if needed based on actual grading scale
        bins = [0, 3.0, 4.0, 5.01] # Example: <3.0 (Excellent), 3.0-4.0 (Good), >4.0 (Alarming/Failing)
        labels = ['<3.0 (Excellent)', '3.0-4.0 (Good)', '>4.0 (Alarming/Failing)']
        relevant_students_df['Grade Category'] = pd.cut(relevant_students_df['Final Grade'], bins=bins, labels=labels, right=False)

        # Create histogram
        fig_hist = px.histogram(relevant_students_df, x='Grade Category',
                              title=f'Grade Distribution for {first_course_code}',
                              labels={'Grade Category': 'Grade Category', 'count': 'Number of Students'},
                              color='Grade Category', color_discrete_sequence=px.colors.qualitative.Set2)
        fig_hist.update_layout(
            yaxis_title='Number of Students',
            xaxis_title='Grade Category',
            showlegend=False
        )
        fig_hist.show()
    else:
        print(f"  No student data available for predicted risk distribution for {first_course_code}.")

    # 10.7. Pass Rates by Gender Bar Chart
    print("📊 7. Pass Rates by Gender")
    if class_profile_path:
        try:
            # Read the original class profile to get the gender data
            df_raw = pd.read_excel(class_profile_path, header=None)
            # Look for the actual header row
            header_row = None
            for i in range(min(10, df_raw.shape[0])):
                row = df_raw.iloc[i].astype(str).str.upper()
                row_text = ' '.join(row.values).upper()
                if any(word in row_text for word in ['PROGRAM', 'COURSE', 'ENROLLED', 'PASSED', 'STUDENTS']):
                    header_row = i
                    break

            if header_row is None:
                header_row = 4

            # Read with proper header
            df_gender = pd.read_excel(class_profile_path, header=header_row)
            # Remove rows that are clearly not data rows
            df_gender = df_gender[~df_gender.iloc[:, 0].astype(str).str.upper().isin(['TOTAL', 'PERCENTAGE', 'PREPARED BY:', 'CHECKED:', 'APPROVED BY:', 'NOTED:'])]

            # Calculate gender-specific pass rates
            if df_gender.shape[1] >= 8:
                # Use pd.to_numeric for gender data as well, handling errors
                total_male_enrolled = pd.to_numeric(df_gender.iloc[:, 2], errors='coerce').sum() if df_gender.shape[1] > 2 else 0
                total_female_enrolled = pd.to_numeric(df_gender.iloc[:, 3], errors='coerce').sum() if df_gender.shape[1] > 3 else 0
                total_male_passed = pd.to_numeric(df_gender.iloc[:, 6], errors='coerce').sum() if df_gender.shape[1] > 6 else 0
                total_female_passed = pd.to_numeric(df_gender.iloc[:, 7], errors='coerce').sum() if df_gender.shape[1] > 7 else 0

                # Calculate pass rates, handling potential division by zero or NaN
                male_pass_rate = (total_male_passed / total_male_enrolled * 100) if total_male_enrolled > 0 and pd.notna(total_male_passed) and pd.notna(total_male_enrolled) else 0
                female_pass_rate = (total_female_passed / total_female_enrolled * 100) if total_female_enrolled > 0 and pd.notna(total_female_passed) and pd.notna(total_female_enrolled) else 0
                overall_pass_rate = (total_male_passed + total_female_passed) / (total_male_enrolled + total_female_enrolled) * 100 if (total_male_enrolled + total_female_enrolled) > 0 and pd.notna(total_male_passed + total_female_passed) and pd.notna(total_male_enrolled + total_female_enrolled) else 0

                # Create bar chart
                gender_pass_data = pd.DataFrame({
                    'Gender': ['Males', 'Females', 'Overall'],
                    'Pass Rate': [male_pass_rate, female_pass_rate, overall_pass_rate]
                })
                fig_gender_pass = px.bar(gender_pass_data, x='Gender', y='Pass Rate',
                                        title='Pass Rate Comparison by Gender',
                                        labels={'Pass Rate': 'Pass Rate (%)', 'Gender': 'Gender'},
                                        color='Gender', color_discrete_sequence=px.colors.qualitative.Set3)
                fig_gender_pass.update_layout(yaxis_range=[0, 100])
                fig_gender_pass.show()
            else:
                print("  Not enough columns found for Pass Rates by Gender chart.")
        except Exception as e:
            print(f"  Error generating Pass Rates by Gender chart: {e}")
    else:
        print("  Class profile not provided for Pass Rates by Gender chart.")

    # 11. === NEW: Precision, Accuracy, F1-Score Calculation ===
    print("="*80)
    print("🔍 PRECISION, ACCURACY, F1-SCORE CALCULATION")
    print("="*80)
    try:
        # Calculate ground truth binary labels (1 for passed, 0 for failed)
        y_true = []
        y_pred = []
        for _, row in df.iterrows():
            # For each subject, create labels for each enrolled student
            # Ground truth: 1 for each passed student, 0 for each failed student
            y_true.extend([1] * int(row['PASSED']))  # Add 'PASSED' number of 1s
            y_true.extend([0] * int(row['FAILED']))  # Add 'FAILED' number of 0s

            # Prediction: Based on the overall pass rate of the subject
            # Let's assume a simple prediction where we predict the class average
            # This is a basic example. A more complex model might predict individual student outcomes.
            # For simplicity here, we'll predict 1 (pass) if the subject's pass rate is >= 75%, else 0 (fail)
            # This creates a prediction for each student based on the subject's overall performance
            threshold = 75.0 # Define a threshold for prediction
            predicted_label = 1 if row['PASS_RATE'] >= threshold else 0
            y_pred.extend([predicted_label] * int(row['ENROLLED'])) # Add 'ENROLLED' number of predictions based on subject pass rate

        # Ensure lengths match (important if data was inconsistent)
        if len(y_true) != len(y_pred):
            print(f"⚠️ Length mismatch between y_true ({len(y_true)}) and y_pred ({len(y_pred)}). Calculating metrics on truncated arrays.")
            min_len = min(len(y_true), len(y_pred))
            y_true = y_true[:min_len]
            y_pred = y_pred[:min_len]

        if len(y_true) == 0:
            print("⚠️ No data points available for metric calculation.")
            return # Exit the function if no data

        # Calculate metrics
        accuracy = accuracy_score(y_true, y_pred)
        # Use 'binary' averaging for precision and recall if only one class is present
        # Otherwise, use 'weighted' to handle potential multi-class scenario (though here it's binary)
        # Also handle cases where labels might be missing in either y_true or y_pred
        unique_labels_true = set(y_true)
        unique_labels_pred = set(y_pred)
        all_labels = sorted(list(unique_labels_true.union(unique_labels_pred)))

        # Calculate precision and recall for each class
        precision_per_class = precision_score(y_true, y_pred, labels=all_labels, average=None, zero_division=0)
        recall_per_class = recall_score(y_true, y_pred, labels=all_labels, average=None, zero_division=0)

        # Calculate weighted average precision and recall
        precision_weighted = precision_score(y_true, y_pred, labels=all_labels, average='weighted', zero_division=0)
        recall_weighted = recall_score(y_true, y_pred, labels=all_labels, average='weighted', zero_division=0)

        # Calculate F1-Score using the weighted averages
        f1_weighted = f1_score(y_true, y_pred, labels=all_labels, average='weighted', zero_division=0)
        f1_macro = f1_score(y_true, y_pred, labels=all_labels, average='macro', zero_division=0)

        # Display results
        print(f"📊 Accuracy: {accuracy:.4f}")
        print(f"📊 Weighted Precision: {precision_weighted:.4f}")
        print(f"📊 Weighted Recall (Sensitivity): {recall_weighted:.4f}")
        print(f"📊 Weighted F1-Score: {f1_weighted:.4f}")
        print(f"📊 Macro F1-Score: {f1_macro:.4f}")
        print("-" * 80)
        print("Class-specific Metrics (if applicable):")
        for i, label in enumerate(all_labels):
            print(f"  Class {label} - Precision: {precision_per_class[i]:.4f}, Recall: {recall_per_class[i]:.4f}")
    except Exception as e:
        print(f"❌ Error calculating Precision, Accuracy, F1-Score: {e}")
        import traceback
        traceback.print_exc() # Print detailed error traceback for debugging


# --- Modified run_analysis function to capture ALL output ---
def run_analysis(class_profile_path, def_report_path, faculty_name, class_record_path_map=None, program=None, semester=None, idempotency_key=None, academic_year=None):
    """Runs the full analysis, captures ALL output, and saves results to the database."""
    print("🔍 Starting analysis for upload...")
    print("📘 Preview: Class Profile (first rows)")
    try:
        _cp = str(class_profile_path).lower()
        if _cp.endswith('.pdf'):
            preview_pdf(class_profile_path)
        elif _cp.endswith('.docx'):
            preview_docx(class_profile_path)
        else:
            preview_excel(class_profile_path)
    except Exception as e:
        print(f"⚠️ Class profile preview error: {e}")
    print("📄 Preview: Deficiency Report (first rows)")
    try:
        if str(def_report_path).lower().endswith('.pdf'):
            preview_pdf(def_report_path)
        else:
            preview_docx(def_report_path)
    except Exception as e:
        print(f"⚠️ Deficiency report preview error: {e}")
    if class_record_path_map: # Added preview for class record
        print("📘 Preview: Class Record Map (course -> file):")
        for path, course in class_record_path_map.items():
             print(f"   {course}: {os.path.basename(path)}")

    analysis_warnings = []
    _cp = str(class_profile_path).lower()
    if _cp.endswith('.csv') or _cp.endswith('.xlsx') or _cp.endswith('.xls'):
        df = read_spreadsheet_any(class_profile_path)
    elif _cp.endswith('.pdf'):
        df = preview_pdf(class_profile_path)
        if df is None:
            analysis_warnings.append("The PDF file could not be processed for Class Profile. Please upload a clearer/standard text-based PDF or a spreadsheet.")
            raise ValueError("Class Profile PDF lacked extractable tables")
    elif _cp.endswith('.docx'):
        df = preview_docx(class_profile_path)
        if df is None:
            analysis_warnings.append("The DOCX file could not be processed for Class Profile. Please upload a document with a clear table or a spreadsheet.")
            raise ValueError("Class Profile DOCX lacked extractable tables")
    else:
        try:
            df = pd.read_excel(class_profile_path)
        except Exception as e:
            analysis_warnings.append("Unsupported Class Profile format; please use CSV/XLSX or a PDF/DOCX with tables.")
            raise
    df.columns = make_unique_columns(df.columns)
    # Remove rows that are clearly not data rows (like "TOTAL", "PERCENTAGE", "Prepared by")
    # Improved filtering: check for exact matches and partial matches
    first_col = df.iloc[:, 0].astype(str).str.upper().str.strip()
    bad_keywords = ['TOTAL', 'PERCENTAGE', 'PREPARED BY:', 'CHECKED:', 'APPROVED BY:', 'NOTED:', 'PROGRAM, YEAR & SECTION']
    
    # Filter out rows that start with or equal bad keywords
    keep_mask = ~first_col.isin(bad_keywords)
    for kw in ['PREPARED BY', 'CHECKED BY', 'APPROVED BY', 'NOTED BY', 'PROGRAM, YEAR & SECTION']:
        keep_mask = keep_mask & ~first_col.str.contains(kw, regex=False)
        
    df = df[keep_mask]

    program_col = smart_find_col(df, ['PROGRAM', 'PROGRAMME', 'COURSE PROGRAM', 'SECTION', 'YEAR & SECTION'])
    
    def _find(words):
        uw = [w.upper() for w in words]
        for c in df.columns:
            uc = c.upper()
            if all(w in uc for w in uw):
                return c
        return None
    cols = list(df.columns)
    def _find_index(words):
        uw = [w.upper() for w in words]
        for i, c in enumerate(cols):
            uc = str(c).upper()
            if all(w in uc for w in uw):
                return i
        return -1
    course_col = smart_find_col(df, ['COURSE', 'SUBJECT', 'COURSE CODE', 'SUBJECT CODE', 'DESCRIPTIVE TITLE', 'TITLE', 'CODE']) or _find(['COURSE'])
    enrolled_total_col = _find(['ENROLLED', 'TOTAL']) or _find(['STUDENTS', 'ENROLLED', 'TOTAL'])
    passed_total_col = _find(['PASSED', 'TOTAL']) or _find(['PASSED STUDENTS', 'TOTAL'])
    enrolled_male = _find(['ENROLLED', 'MALE']) or _find(['STUDENTS', 'ENROLLED', 'MALE'])
    enrolled_female = _find(['ENROLLED', 'FEMALE']) or _find(['STUDENTS', 'ENROLLED', 'FEMALE'])
    passed_male = _find(['PASSED', 'MALE'])
    passed_female = _find(['PASSED', 'FEMALE'])
    i_en = _find_index(['STUDENTS', 'ENROLLED'])
    i_pass = _find_index(['PASSED', 'STUDENTS'])

    inferred = []
    try:
        course_idx = list(df.columns).index(course_col)
        inferred = _infer_group_triplets(df, start_idx=course_idx+1)
    except Exception:
        inferred = _infer_group_triplets(df, start_idx=2)

    # Prefer header-based detection first
    header_en = find_group_columns_by_header(df, 'ENROLLED') or find_group_columns_by_header(df, 'STUDENTS ENROLLED')
    header_pass = find_group_columns_by_header(df, 'PASSED') or find_group_columns_by_header(df, 'PASSED STUDENTS')
    header_failed = find_group_columns_by_header(df, 'FAILED') or find_group_columns_by_header(df, 'FAILED STUDENTS')
    header_incomplete = find_group_columns_by_header(df, 'INCOMPLETE') or find_group_columns_by_header(df, 'INCOMPLETE STUDENTS')
    header_drop = find_group_columns_by_header(df, 'DROP') or find_group_columns_by_header(df, 'DROP STUDENTS')
    header_condfail = find_group_columns_by_header(df, 'CONDITIONAL FAILURE') or find_group_columns_by_header(df, 'CONDITIONAL FAILURE STUDENTS')

    if course_col is None:
        # Fallback: if first column is program, try second column
        if program_col and df.columns[0] == program_col and len(df.columns) > 1:
            course_col = df.columns[1]
            analysis_warnings.append(f"Could not detect COURSE column; inferred second column '{course_col}' as course.")
        else:
            course_col = df.columns[0]
            analysis_warnings.append("Could not detect COURSE column; using first column as course.")

    # --- Content-Based Column Detection (Overrides header detection if high confidence) ---
    best_course_col = None
    best_course_score = 0.0
    best_program_col = None
    best_program_score = 0.0
    
    # Analyze first 20 valid rows
    sample_df = df.head(20).copy()
    
    for col in df.columns:
        # Convert to string and clean
        vals = sample_df[col].dropna().astype(str).str.strip()
        if vals.empty:
            continue
            
        # 1. Subject Code Score
        # Matches patterns like: GEC 105, IT 101, CC 102, ITE-101, MATH 1
        # Logic: 2-5 letters, optional space/dash, 1-4 digits, optional suffix char
        # Must NOT start with common Program prefixes unless it's a specific subject code format
        def is_subject_code(v):
            if re.search(r'^(BS|MS|AB|ACT|BIT)', v, re.IGNORECASE): return False # Likely a program
            return bool(re.search(r'^[A-Z]{2,5}[\s-]?\d{1,4}[A-Z]?$', v, re.IGNORECASE))
            
        subj_matches = vals.apply(is_subject_code).sum()
        subj_score = subj_matches / len(vals)
        
        if subj_score > best_course_score:
            best_course_score = subj_score
            best_course_col = col
            
        # 2. Program/Section Score
        # Matches patterns like: BSIT 1A, BSCS 2B, ACT 1, BIT 3C
        def is_program_section(v):
            return bool(re.search(r'^(BS|MS|AB|ACT|BIT)[A-Z]*[\s-]?\d{1,2}[A-Z]?$', v, re.IGNORECASE)) or \
                   bool(re.search(r'^(BSIT|BSCS|BSIS|ACT|BIT)', v, re.IGNORECASE))

        prog_matches = vals.apply(is_program_section).sum()
        prog_score = prog_matches / len(vals)
        
        if prog_score > best_program_score:
            best_program_score = prog_score
            best_program_col = col

    # Apply overrides if we found strong matches (>40% of rows match pattern)
    if best_course_score > 0.4:
        print(f"🔍 Content-based detection: '{best_course_col}' identified as SUBJECT CODE (score: {best_course_score:.2f})")
        # If the header-based finder picked a different one, or none, use this
        course_col = best_course_col
        
    if best_program_score > 0.4:
        print(f"🔍 Content-based detection: '{best_program_col}' identified as PROGRAM/SECTION (score: {best_program_score:.2f})")
        program_col = best_program_col
        
    # Prevent same column being used for both if scores are close?
    if course_col == program_col and course_col is not None:
        # Tie-breaker: which score is higher?
        if best_course_score > best_program_score:
            program_col = None # Must find another program col?
        else:
            course_col = None

    # --------------------------------------------------------------------------------

    df['COURSE'] = df[course_col].astype(str).apply(clean_str)
    # Drop summary rows that skew rates
    bad_course_keywords = ['TOTAL', 'PERCENTAGE', 'PROGRAM, YEAR & SECTION', 'PREPARED BY', 'CHECKED BY', 'APPROVED BY', 'NOTED BY']
    df = df[~df['COURSE'].str.upper().isin(bad_course_keywords)]
    # Also filter if course starts with any bad keyword
    for kw in bad_course_keywords:
        df = df[~df['COURSE'].str.upper().str.startswith(kw)]

    df['PROGRAM'] = df[program_col].astype(str).apply(clean_str) if program_col else "UNKNOWN"

    # Validation: Check if COURSE and PROGRAM columns appear swapped (common in some PDF layouts)
    try:
        sample_course = df['COURSE'].dropna().head(10).astype(str).tolist()
        sample_program = df['PROGRAM'].dropna().head(10).astype(str).tolist()
        
        # Simple heuristics
        def looks_like_section(vals):
            # Check for common program codes in section names (e.g., BSIT 1A)
            return any(re.search(r'(BSIT|BSCS|BSIS|ACT|BIT)', v, re.IGNORECASE) for v in vals)
            
        def looks_like_subject(vals):
            # Check for subject code pattern (e.g., GEC 106, ITE 301) - 3-4 letters followed by digits
            # Ensure it doesn't also look like a section
            count = 0
            for v in vals:
                 if re.search(r'^[A-Z]{3,4}[\s-]?\d{1,3}', v, re.IGNORECASE) and not re.search(r'(BSIT|BSCS)', v, re.IGNORECASE):
                     count += 1
            return count > 0

        course_is_section = looks_like_section(sample_course)
        program_is_subject = looks_like_subject(sample_program)
        
        # Also check the inverse to be sure
        course_is_subject = looks_like_subject(sample_course)
        program_is_section = looks_like_section(sample_program)
        
        if (course_is_section and program_is_subject) and not (course_is_subject and program_is_section):
            print("⚠️ Detected swapped COURSE and PROGRAM/SECTION columns. Swapping them back.")
            analysis_warnings.append("Detected swapped COURSE and PROGRAM columns based on content; swapped them back.")
            # Swap values by renaming
            df.rename(columns={'COURSE': 'TEMP_COURSE'}, inplace=True)
            df.rename(columns={'PROGRAM': 'COURSE'}, inplace=True)
            df.rename(columns={'TEMP_COURSE': 'PROGRAM'}, inplace=True)
    except Exception as e:
        print(f"⚠️ Error checking for swapped columns: {e}")

    if header_en:
        df['ENROLLED'] = sanitize_series(df.iloc[:, header_en['total_idx']]).astype(int)
    elif i_en >= 0 and (i_en + 3) < len(cols):
        total_series = sanitize_series(df.iloc[:, i_en + 3])
        df['ENROLLED'] = total_series.astype(int)
    elif i_en >= 0 and (i_en + 2) < len(cols):
        em = sanitize_series(df.iloc[:, i_en + 1])
        ef = sanitize_series(df.iloc[:, i_en + 2])
        df['ENROLLED'] = (em + ef).astype(int)
    elif enrolled_total_col:
        df['ENROLLED'] = sanitize_series(df[enrolled_total_col]).astype(int)
    elif enrolled_male or enrolled_female:
        em = sanitize_series(df[enrolled_male] if enrolled_male else 0)
        ef = sanitize_series(df[enrolled_female] if enrolled_female else 0)
        df['ENROLLED'] = (em + ef).astype(int)
    elif inferred:
        s = pd.to_numeric(df.iloc[:, inferred[0]['total_idx']], errors='coerce').fillna(0)
        df['ENROLLED'] = s.astype(int)
        analysis_warnings.append('Mapped ENROLLED totals via heuristic triplet detection.')
    else:
        analysis_warnings.append("Could not detect ENROLLED totals; defaulting to 0.")
        df['ENROLLED'] = 0

    if header_pass:
        df['PASSED'] = sanitize_series(df.iloc[:, header_pass['total_idx']]).astype(int)
    elif i_pass >= 0 and (i_pass + 3) < len(cols):
        total_series = sanitize_series(df.iloc[:, i_pass + 3])
        df['PASSED'] = total_series.astype(int)
    elif i_pass >= 0 and (i_pass + 2) < len(cols):
        pm = sanitize_series(df.iloc[:, i_pass + 1])
        pf = sanitize_series(df.iloc[:, i_pass + 2])
        df['PASSED'] = (pm + pf).astype(int)
    elif passed_total_col:
        df['PASSED'] = sanitize_series(df[passed_total_col]).astype(int)
    elif passed_male or passed_female:
        pm = sanitize_series(df[passed_male] if passed_male else 0)
        pf = sanitize_series(df[passed_female] if passed_female else 0)
        df['PASSED'] = (pm + pf).astype(int)
    elif len(inferred) >= 2:
        s = pd.to_numeric(df.iloc[:, inferred[1]['total_idx']], errors='coerce').fillna(0)
        df['PASSED'] = s.astype(int)
        analysis_warnings.append('Mapped PASSED totals via heuristic triplet detection.')
    else:
        analysis_warnings.append("Could not detect PASSED totals; defaulting to 0.")
        df['PASSED'] = 0
    if header_failed:
        df['FAILED_EXPLICIT'] = sanitize_series(df.iloc[:, header_failed['total_idx']]).astype(int)
    i_fail = _find_index(['FAILED', 'STUDENTS'])
    if 'FAILED_EXPLICIT' not in df.columns and i_fail >= 0 and (i_fail + 2) < len(cols):
        fm = sanitize_series(df.iloc[:, i_fail + 1]) if (i_fail + 1) < len(cols) else 0
        ff = sanitize_series(df.iloc[:, i_fail + 2]) if (i_fail + 2) < len(cols) else 0
        df['FAILED_EXPLICIT'] = (fm + ff).astype(int)
    if header_incomplete:
        df['INCOMPLETE'] = sanitize_series(df.iloc[:, header_incomplete['total_idx']]).astype(int)
    else:
        df['INCOMPLETE'] = 0
    if header_drop:
        df['DROP'] = sanitize_series(df.iloc[:, header_drop['total_idx']]).astype(int)
    else:
        df['DROP'] = 0
    if header_condfail:
        df['COND_FAIL'] = sanitize_series(df.iloc[:, header_condfail['total_idx']]).astype(int)
    else:
        df['COND_FAIL'] = 0
    # If explicit FAILED totals exist, prefer them; also derive residual PASSED when header was missed
    residual_pass = (df['ENROLLED'] - df.get('FAILED_EXPLICIT', 0) - df.get('INCOMPLETE', 0) - df.get('DROP', 0) - df.get('COND_FAIL', 0)).clip(lower=0)
    if not header_pass or (int(pd.to_numeric(df['PASSED'], errors='coerce').sum()) < int(pd.to_numeric(residual_pass, errors='coerce').sum() * 0.3)):
        # Fallback: trust residual pass if current PASSED is suspiciously low
        df['PASSED'] = residual_pass.astype(int)
    df['PASS_RATE'] = np.where(df['ENROLLED'] > 0, (df['PASSED'] / df['ENROLLED'] * 100).round(2), 0)
    computed_failed = (df['ENROLLED'] - df['PASSED'] - df.get('INCOMPLETE', 0) - df.get('DROP', 0) - df.get('COND_FAIL', 0)).clip(lower=0)
    if 'FAILED_EXPLICIT' in df.columns:
        df['FAILED'] = np.where(df['FAILED_EXPLICIT'] > 0, df['FAILED_EXPLICIT'], computed_failed)
    else:
        df['FAILED'] = computed_failed
    df, val_warnings = validate_counts(df)
    analysis_warnings.extend(val_warnings)

    # Deficiency report
    df_def = None
    total_deficiency_students = 0
    try:
        if str(def_report_path).lower().endswith('.pdf'):
            df_def = preview_pdf(def_report_path)
            if df_def is None:
                analysis_warnings.append("The PDF file could not be processed. Please upload a clearer/standard text-based PDF.")
        else:
            df_def = preview_docx(def_report_path)
        if df_def is None:
            raise ValueError("No table detected in deficiency report.")
        # Normalize column names for robust matching
        df_def.columns = make_unique_columns(df_def.columns)
        name_col = smart_find_col(df_def, ['NAME', 'STUDENT NAME', 'FULL NAME', 'STUDENT'])
        course_col2 = smart_find_col(df_def, ['COURSE', 'SUBJECT', 'COURSE CODE', 'SUBJECT CODE', 'TITLE', 'CODE'])
        reason_col = smart_find_col(df_def, ['REASON', 'DEFICIENCY', 'CAUSE', 'STATUS'])

        if course_col2:
            # If we have student names, count UNIQUE deficiency students per course
            if name_col:
                # Clean name values to avoid duplicates due to spacing/case
                df_def[name_col] = df_def[name_col].astype(str).str.strip().str.upper()
                grouped = (
                    df_def.groupby(course_col2)[name_col]
                    .nunique()
                    .reset_index()
                    .rename(columns={course_col2: 'COURSE', name_col: 'NUM_DEF'})
                )
                total_deficiency_students = int(df_def[name_col].nunique())
            # Fallback: count rows per course using reason/status if no name column
            elif reason_col:
                grouped = (
                    df_def.groupby(course_col2)[reason_col]
                    .count()
                    .reset_index()
                    .rename(columns={course_col2: 'COURSE', reason_col: 'NUM_DEF'})
                )
                total_deficiency_students = int(len(df_def.index))
            else:
                # Last resort: just count rows per course
                grouped = (
                    df_def.groupby(course_col2)
                    .size()
                    .reset_index(name='NUM_DEF')
                    .rename(columns={course_col2: 'COURSE'})
                )
                total_deficiency_students = int(len(df_def.index))
            # Normalize grouped course names and ensure numeric NUM_DEF for merging
            if 'COURSE' in grouped.columns:
                grouped['COURSE'] = grouped['COURSE'].astype(str).apply(clean_str)
            if 'NUM_DEF' in grouped.columns:
                grouped['NUM_DEF'] = pd.to_numeric(grouped['NUM_DEF'], errors='coerce').fillna(0).astype(int)
        else:
            print("⚠️ Could not find a course/subject column in deficiency report")
            grouped = pd.DataFrame(columns=['COURSE', 'NUM_DEF'])
            # Attempt to estimate total deficiency students from name column if available
            if name_col:
                total_deficiency_students = int(df_def[name_col].astype(str).str.strip().str.upper().nunique())
            else:
                total_deficiency_students = 0
    except Exception as e:
        print("⚠️ Could not parse deficiency report:", e)
        if str(def_report_path).lower().endswith('.pdf'):
            analysis_warnings.append("The PDF file could not be processed. Please upload a clearer/standard text-based PDF.")
        grouped = pd.DataFrame(columns=['COURSE', 'NUM_DEF'])
        total_deficiency_students = 0

    merged = df.merge(grouped, on='COURSE', how='left').fillna({'NUM_DEF':0})
    # Aggregate duplicate courses within the same upload
    agg_cols = {
        'ENROLLED':'sum', 'PASSED':'sum', 'FAILED':'sum', 'NUM_DEF':'sum'
    }
    if 'INCOMPLETE' in merged.columns:
        agg_cols['INCOMPLETE'] = 'sum'
    if 'DROP' in merged.columns:
        agg_cols['DROP'] = 'sum'
    if 'COND_FAIL' in merged.columns:
        agg_cols['COND_FAIL'] = 'sum'
    agg = merged.groupby(['COURSE', 'PROGRAM'], as_index=False).agg(agg_cols)
    agg['PASS_RATE'] = np.where(agg['ENROLLED']>0, (agg['PASSED']/agg['ENROLLED']*100).round(2), 0)
    merged = agg

    # Pass both course name and program name to the recommendation function
    merged['RECOMMENDATION'] = merged.apply(lambda row: generate_recommendation(row, row['COURSE'], row['PROGRAM']), axis=1)
    merged['FACULTY'] = faculty_name
    if academic_year:
        merged['ACADEMIC_YEAR'] = academic_year
    merged['ANALYSIS_DATE'] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    # Remove rows with NaN in key columns and clean up
    merged = merged.dropna(subset=['COURSE', 'ENROLLED', 'PASSED'])
    merged = merged[merged['COURSE'] != 'nan']  # Remove rows where course is literally 'nan'
    merged = merged[merged['COURSE'] != '']     # Remove empty course names

    # Display and save
    print(f"<h3>AI Analysis Summary for <em>{faculty_name}</em></h3>")
    print(merged.head(20)) # Use print instead of display for backend
    out_dir = os.path.join(RESULTS_DIR, faculty_name)
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"{faculty_name}_AI_Report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv")
    merged.to_csv(out_path, index=False)
    print(f"✅ Analysis complete. Saved → {out_path}")

    master_df = load_master()
    master_df = pd.concat([master_df, merged], ignore_index=True)
    save_master(master_df)

    # --- NEW: Capture ALL output from generate_student_analytics ---
    detailed_output_str = io.StringIO()
    error_buffer = io.StringIO()
    try:
        with redirect_stdout(detailed_output_str), redirect_stderr(error_buffer):
            generate_student_analytics(merged, faculty_name, class_profile_path, def_report_path, class_record_path_map, academic_year=academic_year)
        detailed_output_str = detailed_output_str.getvalue()
        error_str = error_buffer.getvalue()
        if error_str:
            print(f"generate_student_analytics generated errors: {error_str}") # Log errors
    except Exception as e:
        print(f"Error capturing output from generate_student_analytics: {e}")
        import traceback
        traceback.print_exc()
        detailed_output_str = f"Error capturing output: {e}"


    # Show plots
    try:
        fig = go.Figure()
        fig.add_trace(go.Bar(x=merged['COURSE'], y=merged['PASS_RATE'], name='Pass Rate'))
        fig.update_layout(title=f"Pass Rate Summary – {faculty_name}", height=400)
        fig.show()

        # --- TEXT-BASED RECOMMENDATION SUMMARY ---
        print("="*80)
        print("📋 RECOMMENDATION SUMMARY")
        print("="*80)
        # Count occurrences of each recommendation
        rec_counts = Counter(merged['RECOMMENDATION'])
        print("Most Common Recommendations:")
        print("-" * 80)
        for rec, count in rec_counts.most_common(): # Iterate through sorted recommendations
            print(f"• ({count}) {rec}")
        print("-" * 80)

        # --- NEW: Summarized Recommendations by Type ---
        print("📋 RECOMMENDATION SUMMARY (Grouped by Type):")
        print("="*80)
        # Dictionary to hold counts for each type of recommendation
        type_counts = Counter()
        for full_recommendation in merged['RECOMMENDATION']:
            parts = full_recommendation.split(" | ")
            for part in parts:
                part = part.strip()
                # Identify the type based on prefixes or keywords
                if part.startswith("🎉 Excellent performance"):
                    type_counts["Performance - Excellent"] += 1
                elif part.startswith("👍 Good performance"):
                    type_counts["Performance - Good"] += 1
                elif part.startswith("⚠️ Fair performance"):
                    type_counts["Performance - Fair"] += 1
                elif part.startswith("🚨 Poor performance"):
                    type_counts["Performance - Poor"] += 1
                elif part.startswith("📋 Address"):
                    type_counts["Deficiency Management"] += 1
                elif part.startswith("📈 Strong employment outlook"):
                    type_counts["Employment Outlook - Strong"] += 1
                elif part.startswith("📈 Good employment outlook"):
                    type_counts["Employment Outlook - Good"] += 1
                elif part.startswith("📈 Employment outlook"):
                    type_counts["Employment Outlook - Other"] += 1
                elif part.startswith("💼 Potential career paths"):
                    type_counts["Job Path Alignment"] += 1
                elif part.startswith("💡 LSPU Suggestion"):
                    type_counts["LSPU Improvement Suggestion"] += 1
                elif part.startswith("⚠️ LSPU Data"):
                    type_counts["LSPU Deficiency Reason"] += 1
                elif part.startswith("🎓 Student Suggestion"):
                    type_counts["Student Suggestion"] += 1
                elif part.startswith("🔍 Internal Review Weakness"):
                    type_counts["Internal Review Weakness"] += 1
                elif part.startswith("✅ Internal Review Action"):
                    type_counts["Internal Review Action"] += 1
                elif part.startswith("📋 Internal Review Recommendation"):
                    type_counts["Internal Review Recommendation"] += 1
                elif any(kw in part.lower() for kw in ["large class", "peer mentoring", "teaching assistants"]):
                    type_counts["Class Size Management"] += 1
                elif any(kw in part.lower() for kw in ["small class", "personalized attention"]):
                    type_counts["Small Class Benefit"] += 1
                else:
                    # Group other generic recommendations
                    type_counts["Generic/Subject-Specific"] += 1

        # Print the grouped counts
        for rec_type, count in type_counts.most_common():
            print(f"• {rec_type}: {count} occurrences")
        # ----------------------------------------------------------

    except Exception as e:
        print("⚠️ Visualization skipped:", e)

    # --- Create Database Entry for Faculty Upload ---
    # Store metadata in file_paths JSON to avoid schema migrations
    try:
        meta = {
            'files': {
                'class_profile': class_profile_path,
                'def_report': def_report_path,
                'class_records': list(class_record_path_map.keys()) if class_record_path_map else []
            },
            'metadata': {
                'program': program or '',
                'semester': semester or '',
                'academic_year': academic_year or '',
                'idempotency_key': idempotency_key or ''
            }
        }
        file_paths_json = json.dumps(meta)
    except Exception:
        file_paths_json = None

    upload_entry = FacultyUpload(faculty_name=faculty_name, file_paths=file_paths_json, academic_year=academic_year)
    db.session.add(upload_entry)
    db.session.flush() # Get the ID for the new upload

    # --- Process each subject and save to DB ---
    for _, row in merged.iterrows():
        subject_entry = SubjectAnalysis(
            upload_id=upload_entry.id,
            course=row['COURSE'],
            program=row['PROGRAM'],
            enrolled=row['ENROLLED'],
            passed=row['PASSED'],
            failed=row['FAILED'],
            pass_rate=row['PASS_RATE'],
            num_def=row['NUM_DEF'],
            recommendation=row['RECOMMENDATION']
        )
        db.session.add(subject_entry)

    # --- Save the captured detailed analytics output ---
    analytics_output_entry = AnalyticsData(
        upload_id=upload_entry.id,
        data_type='detailed_analytics_output',
        data_content=detailed_output_str # Store the full output string
    )
    db.session.add(analytics_output_entry)

    # --- Also store a compact summary count of deficiency students ---
    try:
        deficiency_summary_entry = AnalyticsData(
            upload_id=upload_entry.id,
            data_type='deficiency_students_summary',
            data_content=json.dumps({'total_deficiency_students': int(total_deficiency_students)})
        )
        db.session.add(deficiency_summary_entry)
    except Exception as e:
        print(f"⚠️ Could not store deficiency_students_summary: {e}")

    # --- Calculate Analytics Metrics and Save ---
    y_true = []
    y_pred = []
    for _, row in merged.iterrows():
        y_true.extend([1] * int(row['PASSED']))
        y_true.extend([0] * int(row['FAILED']))
        threshold = 75.0
        predicted_label = 1 if row['PASS_RATE'] >= threshold else 0
        y_pred.extend([predicted_label] * int(row['ENROLLED']))

    if len(y_true) != len(y_pred):
        min_len = min(len(y_true), len(y_pred))
        y_true = y_true[:min_len]
        y_pred = y_pred[:min_len]

    if len(y_true) > 0:
        accuracy = accuracy_score(y_true, y_pred)
        precision = precision_score(y_true, y_pred, average='weighted', zero_division=0)
        recall = recall_score(y_true, y_pred, average='weighted', zero_division=0)
        f1 = f1_score(y_true, y_pred, average='weighted', zero_division=0)
        macro_f1 = f1_score(y_true, y_pred, average='macro', zero_division=0)

        # Cast to native Python types to avoid JSON serialization issues (e.g., numpy.float64)
        metrics_data = {
            'accuracy': float(accuracy),
            'precision': float(precision),
            'recall': float(recall),
            'f1_weighted': float(f1),
            'f1_macro': float(macro_f1)
        }
        metrics_entry = AnalyticsData(
            upload_id=upload_entry.id,
            data_type='precision_recall_f1',
            data_content=json.dumps(metrics_data)
        )
        db.session.add(metrics_entry)

    # --- Commit all changes to the database ---
    db.session.commit()
    print(f"✅ Analysis complete and saved to database for {faculty_name}.")

    # Return summary data for the frontend
    # Build extraction report
    extraction_report = {
        'class_profile': {
            'file': os.path.basename(class_profile_path),
            'rows_parsed': int(df.shape[0]),
            'columns_parsed': int(df.shape[1]),
            'warnings': analysis_warnings,
        },
        'deficiency_report': {
            'file': os.path.basename(def_report_path),
            'rows_parsed': int(grouped.shape[0] if 'grouped' in locals() and isinstance(grouped, pd.DataFrame) else 0),
            'columns_parsed': int(grouped.shape[1] if 'grouped' in locals() and isinstance(grouped, pd.DataFrame) else 0),
        }
    }

    return {
        'upload_id': upload_entry.id,
        'faculty_name': faculty_name,
        'analysis_date': to_manila_iso(upload_entry.analysis_date),
        'subjects_processed': len(merged),
        # Cast sums to native ints to avoid numpy.int64 in JSON
        'total_enrolled': int(merged['ENROLLED'].sum()),
        'total_passed': int(merged['PASSED'].sum()),
        'total_failed': int(merged['FAILED'].sum()),
        'overall_pass_rate': float((merged['PASSED'].sum() / merged['ENROLLED'].sum() * 100) if merged['ENROLLED'].sum() > 0 else 0.0),
        'total_deficiencies': int(merged['NUM_DEF'].sum()),
        'metrics': metrics_data,
        'warnings': analysis_warnings,
        'extraction_report': extraction_report
    }
def validate_counts(df):
    warnings = []
    # Basic sanity checks
    total_students = int(df['ENROLLED'].sum() if 'ENROLLED' in df.columns else 0)
    total_passed = int(df['PASSED'].sum() if 'PASSED' in df.columns else 0)
    total_failed = int(df['FAILED'].sum() if 'FAILED' in df.columns else 0)
    if total_students < 0:
        warnings.append('Total students computed as negative; clamped to 0.')
    if total_passed < 0 or total_failed < 0:
        warnings.append('Negative pass/fail counts detected; clamped to 0.')
    # Per-row checks
    if 'ENROLLED' in df.columns and 'PASSED' in df.columns:
        bad_rows = []
        for i, row in df.iterrows():
            e = int(row.get('ENROLLED', 0) or 0)
            p = int(row.get('PASSED', 0) or 0)
            if p > e:
                bad_rows.append(i)
        if bad_rows:
            warnings.append(f"{len(bad_rows)} rows where PASSED > ENROLLED; clamped.")
            for i in bad_rows:
                df.loc[i, 'PASSED'] = int(df.loc[i, 'ENROLLED'])
    if 'ENROLLED' in df.columns and 'PASSED' in df.columns:
        computed_failed = (df['ENROLLED'] - df['PASSED'] - df.get('INCOMPLETE', 0) - df.get('DROP', 0) - df.get('COND_FAIL', 0)).clip(lower=0)
        if 'FAILED_EXPLICIT' in df.columns:
            df['FAILED'] = np.where(df['FAILED_EXPLICIT'] > 0, df['FAILED_EXPLICIT'], computed_failed)
        else:
            df['FAILED'] = computed_failed
        df['PASS_RATE'] = np.where(df['ENROLLED'] > 0, (df['PASSED'] / df['ENROLLED'] * 100).clip(lower=0, upper=100).round(2), 0)
    return df, warnings

def detect_header_tokens(df_raw, tokens, max_rows=20):
    header_row = None
    for i in range(min(max_rows, df_raw.shape[0])):
        row = df_raw.iloc[i].astype(str).str.upper()
        row_text = ' '.join(row.values).upper()
        if any(tok in row_text for tok in tokens):
            header_row = i
            break
    return header_row

def read_spreadsheet_any(path):
    # Try CSV first
    if str(path).lower().endswith('.csv'):
        df = pd.read_csv(path)
        return df
    # Try all sheets for Excel
    def _combine_multi_header_columns(df):
        if isinstance(df.columns, pd.MultiIndex):
            def join_levels(tup):
                parts = [clean_str(x) for x in tup if str(x).strip() != '' and str(x).lower() != 'nan']
                return ' '.join([p for p in parts if p]).strip()
            df.columns = [join_levels(c) for c in df.columns]
        df.columns = make_unique_columns(df.columns)
        return df
    frames = []
    try:
        xls = pd.ExcelFile(path)
        for sheet in xls.sheet_names:
            try:
                df_raw = pd.read_excel(path, sheet_name=sheet, header=None)
                header_row = detect_header_tokens(df_raw, ['PROGRAM', 'COURSE', 'ENROLLED', 'PASSED', 'STUDENTS', 'SUBJECT', 'TITLE', 'CODE', 'YEAR & SECTION'])
                if header_row is None:
                    header_row = 0
                # Try multi-row header to capture group + subheaders
                try:
                    df = pd.read_excel(path, sheet_name=sheet, header=[header_row, header_row+1])
                    df = _combine_multi_header_columns(df)
                except Exception:
                    df = pd.read_excel(path, sheet_name=sheet, header=header_row)
                    df = _combine_multi_header_columns(df)
                # Filter out empty/non-data frames early
                if df.shape[1] >= 3 and df.shape[0] >= 3:
                    frames.append(df)
            except Exception:
                continue
        if frames:
            # Concatenate all sheets vertically, aligning columns by name
            combined = pd.concat(frames, ignore_index=True, sort=False)
            combined.columns = make_unique_columns(combined.columns)
            return combined
    except Exception:
        pass
    # Fallback: generic read_excel
    return pd.read_excel(path)
MANILA_TZ = timezone(timedelta(hours=8))
UTC_TZ = timezone.utc

def to_manila_iso(dt: datetime):
    if dt is None:
        return datetime.now(MANILA_TZ).isoformat()
    base = dt
    if base.tzinfo is None:
        base = base.replace(tzinfo=UTC_TZ)
    return base.astimezone(MANILA_TZ).isoformat()
